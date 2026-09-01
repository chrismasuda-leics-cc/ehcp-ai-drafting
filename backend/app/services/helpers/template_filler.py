import os
import re
import json
import time
import argparse
from copy import deepcopy
from datetime import datetime
from typing import Dict, Any, Optional

try:
    import mammoth
    import openpyxl
    from markdownify import markdownify as mdify
    from bs4 import BeautifulSoup
    from docx import Document
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor
except Exception as e:
    raise ImportError(
        "Missing dependency. Install requirements from requirements.txt: pip install -r requirements.txt"
    ) from e


PLACEHOLDER_PATTERNS = [
    lambda k: re.compile(r"\{\{\s*%s\s*\}\}" % re.escape(k)),
    lambda k: re.compile(r"\[\[\s*%s\s*\]\]" % re.escape(k)),
    lambda k: re.compile(r"<<\s*%s\s*>>" % re.escape(k)),
    lambda k: re.compile(r"%s" % re.escape(k)),
    lambda k: re.compile(r"%s" % re.escape(k), re.IGNORECASE),
]

DEFAULT_FONT_NAME = "Century Gothic"
DEFAULT_FONT_SIZE = 12


def generate_key_variants(key: str):
    """Return a list of plausible placeholder variants for a JSON key.

    Examples: 'preferred_name' -> ['preferred_name', 'preferred name', 'Preferred Name', 'Preferred name', 'preferredname', 'My Preferred Name']
    """
    variants = set()
    key = key.strip()
    variants.add(key)

    # Underscore -> space
    space = key.replace("_", " ")
    variants.add(space)

    # Title and capitalized
    variants.add(space.title())
    variants.add(space.capitalize())
    variants.add(space.lower())
    variants.add(space.upper())

    # Without separators
    variants.add(key.replace("_", ""))

    # With "My " prefix (common label)
    variants.add("My " + space.title())
    variants.add("My " + space)

    # With colon suffix variation (match 'Name:' in docs)
    out = set()
    for v in variants:
        out.add(v)
        out.add(v + ":")
        out.add(v + " :")

    # sort by length desc so longer variants replaced first
    return sorted(out, key=lambda s: -len(s))


def docx_to_markdown(docx_path: str, md_path: str, metadata_path: str = None) -> None:
    """Convert a .docx file to markdown and extract simple metadata about headings/tables/images.

    Saves markdown to `md_path`. If `metadata_path` provided saves JSON metadata there.
    """
    with open(docx_path, "rb") as f:
        result = mammoth.convert_to_html(f)
        html = result.value

    md = mdify(html, heading_style="ATX")
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(md)

    if metadata_path:
        meta = extract_metadata_from_html(html)
        with open(metadata_path, "w", encoding="utf-8") as mf:
            json.dump(meta, mf, indent=2, ensure_ascii=False)


def extract_metadata_from_html(html: str) -> Dict[str, Any]:
    soup = BeautifulSoup(html, "html.parser")
    headings = [h.get_text(strip=True) for h in soup.find_all(re.compile("^h[1-6]$"))]
    tables = len(soup.find_all("table"))
    images = [img.get("src") for img in soup.find_all("img")]
    bold = len(soup.find_all("strong")) + len(soup.find_all("b"))
    italic = len(soup.find_all("em")) + len(soup.find_all("i"))
    lists = len(soup.find_all(["ul", "ol"]))

    return {
        "headings": headings,
        "tables_count": tables,
        "images": images,
        "bold_count": bold,
        "italic_count": italic,
        "lists_count": lists,
    }


def load_json(json_path: str) -> Dict[str, Any]:
    with open(json_path, "r", encoding="utf-8") as f:
        return json.load(f)


def fill_markdown(md_in: str, md_out: str, data: Dict[str, Any]) -> None:
    text = open(md_in, encoding="utf-8").read()

    for key in sorted(data.keys(), key=lambda k: -len(k)):
        value = "" if data[key] is None else str(data[key])
        variants = generate_key_variants(key)
        for var in variants:
            pats = [r"\{\{\s*%s\s*\}\}", r"\[\[\s*%s\s*\]\]", r"<<\s*%s\s*>>"]
            for p in pats:
                text = re.sub(p % re.escape(var), value, text, flags=re.IGNORECASE)
            text = re.sub(r"\b%s\b" % re.escape(var), value, text, flags=re.IGNORECASE)

    with open(md_out, "w", encoding="utf-8") as f:
        f.write(text)


def replace_in_run_text(run_text: str, data: Dict[str, Any]) -> str:
    new_text = run_text
    for key in sorted(data.keys(), key=lambda k: -len(k)):
        value = "" if data[key] is None else str(data[key])
        variants = generate_key_variants(key)
        for var in variants:
            pats = [r"\{\{\s*%s\s*\}\}", r"\[\[\s*%s\s*\]\]", r"<<\s*%s\s*>>"]
            for p in pats:
                new_text = re.sub(p % re.escape(var), value, new_text, flags=re.IGNORECASE)
            new_text = re.sub(r"\b%s\b" % re.escape(var), value, new_text, flags=re.IGNORECASE)
    return new_text


def replace_text_whole(text: str, data: Dict[str, Any]) -> str:
    """Replace placeholders in a larger text block using data variants."""
    new = text
    for key in sorted(data.keys(), key=lambda k: -len(k)):
        value = "" if data[key] is None else str(data[key])
        variants = generate_key_variants(key)
        for var in variants:
            pats = [r"\{\{\s*%s\s*\}\}", r"\[\[\s*%s\s*\]\]", r"<<\s*%s\s*>>"]
            for p in pats:
                new = re.sub(p % re.escape(var), value, new, flags=re.IGNORECASE)
            new = re.sub(r"\b%s\b" % re.escape(var), value, new, flags=re.IGNORECASE)
    return new


def replace_placeholders_in_document(docx_in: str, docx_out: str, data: Dict[str, Any]) -> None:
    doc = Document(docx_in)

    for para in doc.paragraphs:
        para_text = para.text
        new_text = replace_text_whole(para_text, data)
        if new_text != para_text:
            if para.runs:
                para.runs[0].text = new_text
                for r in para.runs[1:]:
                    r.text = ""
            else:
                para.add_run(new_text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para_text = para.text
                    new_text = replace_text_whole(para_text, data)
                    if new_text != para_text:
                        if para.runs:
                            para.runs[0].text = new_text
                            for r in para.runs[1:]:
                                r.text = ""
                        else:
                            para.add_run(new_text)

    for section in doc.sections:
        header = section.header
        for para in header.paragraphs:
            para_text = para.text
            new_text = replace_text_whole(para_text, data)
            if new_text != para_text:
                if para.runs:
                    para.runs[0].text = new_text
                    for r in para.runs[1:]:
                        r.text = ""
                else:
                    para.add_run(new_text)
        footer = section.footer
        for para in footer.paragraphs:
            para_text = para.text
            new_text = replace_text_whole(para_text, data)
            if new_text != para_text:
                if para.runs:
                    para.runs[0].text = new_text
                    for r in para.runs[1:]:
                        r.text = ""
                else:
                    para.add_run(new_text)

    doc.save(docx_out)


STRUCTURED_MAPPING = [
    [("My Name", "name")],
    [("Name", "name"), ("Preferred name", "preferred_name")],
    [("Date of Birth", "date_of_birth")],
    [("Sex", "sex"), ("Ethnicity", "ethnicity")],
    [
        ("Child or young person first language or method of communication", "child_or_young_person_first_language_or_method_of_communication"),
        ("Parent or carer first language or method of communication", "parent_or_carer_first_language_or_method_of_communication"),
    ],
    [
        ("Person with parental responsibility - Parent or carer 1 (or young person if they are applying independently)", "main_contact_parent_or_carer_1"),
        ("Relationship", "main_contact_relationship"),
    ],
    [("Email", "main_contact_email"), ("Telephone number", "main_contact_telephone_number")],
    [("Home Address", "main_contact_home_address")],
    [("Person with parental responsibility - Parent or carer 2", "parent_or_carer_2"), ("Relationship", "parent_or_carer_2_relationship")],
    [("Email", "parent_or_carer_2_email"), ("Telephone number", "parent_or_carer_2_telephone_number")],
    [("Home address (if different to parent or carer 1)", "home_address_if_different_to_parent_or_carer_1")],
    [("Education setting", "education_setting")],
    [("GP and key health professional (where applicable)", "gp_and_key_health_professional")],
    [("NHS Number (if available)", "nhs_number_if_available")],
    [("Known to social care", "known_to_social_care")],
    [("Other professionals involved with this child/ young person family", "other_professionals_involved_with_this_child_young_person_family")],
    [("Currently a looked after child?", "currently_a_looked_after_child")],
    [("Has been a looked after child?", "has_been_a_looked_after_child")],
    [("Is there is a one page profile received (append)?", "is_there_a_one_page_profile_received_append")],
    [("Has child or young person moved into Leicestershire with an existing EHCP?", "has_child_or_young_person_moved_into_leicestershire_with_an_existing_ehcp"), ("Date of Move In", "date_of_move_in")],
]


PERSONAL_DETAILS_SHEET = "PersonalDetails"
EDUCATION_DETAILS_SHEET = "EducationAdvice"
HEALTH_DETAILS_SHEET = "HealthAdvice"
SOCIAL_CARE_DETAILS_SHEET = "SocialCareAdvice"


EDUCATION_SECTION_TO_JSON = {
    "A1 Summary of the child or young person's history": "Parent_Carer_Views",
    "A2 - Summary of the views, interests and aspirations of the child and their parent, or of the young person": None,
    "Summary of Involvement": "Summary_of_Involvement",
    "Cognition and learning": "Cognition_and_Learning",
    "Communication and interaction": "Communication_and_Interaction",
    "Social, emotional and mental health difficulties": "Social_Emotional_Mental_Health",
    "Sensory and/or physical needs": "Sensory_and_Physical",
}

A1_TEMPLATE_SECTION = "A1 Summary of the child or young person's history"
A2_TEMPLATE_SECTION = "A2 - Summary of the views, interests and aspirations of the child and their parent, or of the young person"
SUMMARY_TEMPLATE_SECTION = "Primary area of need"
SEND_TEMPLATE_SECTIONS = {
    "Cognition and learning",
    "Communication and interaction",
    "Social, emotional and mental health difficulties",
    "Sensory and/or physical needs",
}
K_TABLE_HEADING = "K - The advice and information gathered during the EHC needs assessment"
SEND_TABLE_MIN_ROWS = 8


def format_value(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, list):
        parts = [format_value(item).strip() for item in value]
        return "\n".join(part for part in parts if part)
    if isinstance(value, dict):
        parts = []
        for key, item in value.items():
            item_text = format_value(item).strip()
            if item_text:
                parts.append(f"{key}: {item_text}")
        return "\n".join(parts)
    return str(value)


def value_for(data: Dict[str, Any], key: str) -> str:
    value = data.get(key)
    return format_value(value)


def load_personal_mapping(xlsx_path: str, sheet_name: str = PERSONAL_DETAILS_SHEET):
    if not xlsx_path or not os.path.exists(xlsx_path):
        return STRUCTURED_MAPPING

    workbook = openpyxl.load_workbook(xlsx_path, data_only=True)
    if sheet_name not in workbook.sheetnames:
        return STRUCTURED_MAPPING

    sheet = workbook[sheet_name]
    groups = {}
    for index, row in enumerate(sheet.iter_rows(values_only=True)):
        if index == 0:
            continue
        row_group, column_order, label, json_key = row[:4]
        if not label or not json_key:
            continue

        try:
            group_key = int(row_group)
        except Exception:
            group_key = index

        try:
            order_key = int(column_order)
        except Exception:
            order_key = 1

        groups.setdefault(group_key, []).append((order_key, str(label).strip(), str(json_key).strip()))

    if not groups:
        return STRUCTURED_MAPPING

    mapping = []
    for group_key in sorted(groups.keys()):
        ordered = sorted(groups[group_key], key=lambda item: item[0])
        mapping.append([(label, json_key) for _, label, json_key in ordered])
    return mapping


def choose_sheet_name(workbook, preferred_name: str, fallback_names):
    if preferred_name in workbook.sheetnames:
        return preferred_name
    for name in fallback_names:
        if name in workbook.sheetnames:
            return name
    return workbook.active.title


def create_structured_markdown(md_out: str, data: Dict[str, Any], mapping=STRUCTURED_MAPPING) -> None:
    lines = []
    for row in mapping:
        cells = []
        for label, jkey in row:
            val = value_for(data, jkey)
            cells.append(label)
            cells.append(val)
        line = "| " + " | ".join(cells) + " |"
        lines.append(line)

    with open(md_out, "w", encoding="utf-8") as f:
        for l in lines:
            f.write(l + "\n")


def build_augmented_data(data: Dict[str, Any], mapping=STRUCTURED_MAPPING) -> Dict[str, Any]:
    out = dict(data)
    for row in mapping:
        for label, jkey in row:
            val = "" if data.get(jkey) is None else data.get(jkey)
            out[label] = val
            out[label + ":"] = val
            out[label.rstrip(":").strip()] = val
    return out


def normalize_label(text: str) -> str:
    cleaned = text.replace("\n", " ").lower()
    cleaned = cleaned.replace("\u2019", "'")
    cleaned = re.sub(r"[^a-z0-9\s]", " ", cleaned)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    return cleaned


def row_matches_mapping(row, mapping_row) -> bool:
    row_texts = [normalize_label(cell.text) for cell in row.cells]
    return all(any(normalize_label(label) in cell_text for cell_text in row_texts) for label, _ in mapping_row)


def fill_row_values(row, mapping_row, data: Dict[str, Any]) -> None:
    row_texts = [normalize_label(cell.text) for cell in row.cells]
    used_indexes = set()

    for label, json_key in mapping_row:
        label_norm = normalize_label(label)
        match_index = None
        for index, cell_text in enumerate(row_texts):
            if index in used_indexes:
                continue
            if label_norm in cell_text:
                match_index = index
                used_indexes.add(index)
                break

        if match_index is None:
            continue

        value_index = match_index + 1
        if value_index < len(row.cells):
            value = value_for(data, json_key)
            if "email" in json_key.lower():
                set_cell_email_hyperlink(row.cells[value_index], value)
            elif "date" in json_key.lower():
                set_cell_text(row.cells[value_index], format_k_date(value))
            elif label_norm == normalize_label("My Name:"):
                # Hardcoded font size for the name field only; all other fields
                # use get_cell_font_size / apply_font_size_to_cell to respect the template.
                set_cell_text_with_style(row.cells[value_index], value, color_rgb="0F4761", underline=False, font_size=16)
            else:
                set_cell_text(row.cells[value_index], value)


def create_filled_template_docx(template_docx: str, docx_out: str, data: Dict[str, Any], mapping=STRUCTURED_MAPPING) -> None:
    doc = Document(template_docx)
    mapping_index = 0

    for table in doc.tables:
        for row in table.rows:
            if mapping_index < len(mapping):
                current_mapping_row = mapping[mapping_index]
                if row_matches_mapping(row, current_mapping_row):
                    fill_row_values(row, current_mapping_row, data)
                    mapping_index += 1

    apply_font_to_document(doc)
    doc.save(docx_out)


def load_education_mapping(xlsx_path: str, sheet_name: str = EDUCATION_DETAILS_SHEET):
    workbook = openpyxl.load_workbook(xlsx_path, data_only=True)
    actual_sheet_name = choose_sheet_name(workbook, sheet_name, ["EducationDetails", "Sheet1"])
    sheet = workbook[actual_sheet_name]
    rows = []
    for index, row in enumerate(sheet.iter_rows(values_only=True)):
        if index == 0:
            continue
        template_column, template_section, key_field, section_name, template_subcolumn = row[:5]
        if not any([template_column, template_section, key_field, section_name, template_subcolumn]):
            continue
        matching_template_section = next(
            (
                known_section
                for known_section in EDUCATION_SECTION_TO_JSON
                if normalize_label(template_section or "") == normalize_label(known_section)
            ),
            None,
        )
        if matching_template_section and normalize_label(matching_template_section) != normalize_label(A2_TEMPLATE_SECTION):
            expected_section = EDUCATION_SECTION_TO_JSON[matching_template_section]
            if expected_section:
                section_name = expected_section
        rows.append(
            {
                "template_column": (template_column or "").strip() if isinstance(template_column, str) else template_column,
                "template_section": (template_section or "").strip() if isinstance(template_section, str) else template_section,
                "key": (key_field or "").strip() if isinstance(key_field, str) else key_field,
                "section": (section_name or "").strip() if isinstance(section_name, str) else section_name,
                "template_subcolumn": (template_subcolumn or "").strip() if isinstance(template_subcolumn, str) else template_subcolumn,
            }
        )
    return rows


def load_optional_json(json_path: str) -> Dict[str, Any]:
    with open(json_path, "r", encoding="utf-8") as file_handle:
        return json.load(file_handle)


def combine_mapping_values(education_data: Dict[str, Any], section_name: str, key_field: str) -> str:
    if not key_field:
        return ""

    if section_name:
        source_section = education_data.get(section_name, {})
        if not isinstance(source_section, dict):
            return ""
    else:
        source_section = education_data
        if not isinstance(source_section, dict):
            return ""

    values = []
    for key_name in [item.strip() for item in key_field.split(",")]:
        if not key_name:
            continue
        value = source_section.get(key_name)
        if value is None:
            continue
        text = format_value(value).strip()
        if text:
            values.append(text)
    return "\n\n".join(values)


def resolve_mapping_value(data: Dict[str, Any], section_name: str, key_field: str):
    if not key_field:
        return ""

    if section_name:
        source_section = data.get(section_name, {})
        if not isinstance(source_section, dict):
            return ""
    else:
        source_section = data
        if not isinstance(source_section, dict):
            return ""

    key_names = [item.strip() for item in key_field.split(",") if item.strip()]
    if len(key_names) == 1:
        return source_section.get(key_names[0], "")
    return combine_mapping_values(data, section_name, key_field)


def child_perspective_label(personal_data: Dict[str, Any], fallback_label: str) -> str:
    label_norm = normalize_label(fallback_label)
    if "perspective" in label_norm or "voice" in label_norm:
        if "parent" in label_norm or "carer" in label_norm:
            return "Parent/Carer's Voice:"
        return "Child/Young Person's Voice:"
    return fallback_label


def find_table_by_heading(doc: Document, heading_text: str, min_rows: int = 0):
    target = normalize_label(heading_text)
    for table in doc.tables:
        if not table.rows or len(table.rows) < min_rows:
            continue

        first_row = table.rows[0]
        first_joined = " ".join(normalize_label(cell.text) for cell in first_row.cells if cell.text.strip())
        if target and target in first_joined:
            return table

        if len(table.rows) > 1:
            second_row = table.rows[1]
            second_cells = [normalize_label(cell.text) for cell in second_row.cells if cell.text.strip()]
            second_joined = " ".join(second_cells)
            is_header_like = bool(second_cells) and all(len(cell_text.split()) <= 12 for cell_text in second_cells)
            if target and is_header_like and target in second_joined:
                return table
    return None


def get_cell_font_size(cell) -> Optional[float]:
    """Read the font size in points from the first available run in a cell.

    Must be called before clear_cell_content so the template's font size can be
    reapplied to all new runs after the cell content has been written.
    """
    if not cell.paragraphs:
        return None
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if run.font.size is not None:
                return run.font.size.pt
    return None


def apply_font_size_to_cell(cell, font_size_pt: Optional[float]) -> None:
    """Apply a font size in points to every run in a cell.

    Called after writing new content to restore the font size that was captured
    from the template before the cell was cleared.
    """
    if font_size_pt is None:
        return
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(font_size_pt)


def set_cell_text(cell, text: str) -> None:
    value = "" if text is None else str(text)

    if not cell.paragraphs:
        cell.text = value
        return

    first_paragraph = cell.paragraphs[0]
    if first_paragraph.runs:
        first_paragraph.runs[0].text = value
        for run in first_paragraph.runs[1:]:
            run.text = ""
    else:
        first_paragraph.add_run(value)

    for paragraph in cell.paragraphs[1:]:
        for run in paragraph.runs:
            run.text = ""


def set_run_font(run, color_rgb: Optional[str] = None, underline: Optional[bool] = None, font_size: Optional[float] = None) -> None:
    run.font.name = DEFAULT_FONT_NAME
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), DEFAULT_FONT_NAME)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), DEFAULT_FONT_NAME)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:cs"), DEFAULT_FONT_NAME)
    if color_rgb:
        run.font.color.rgb = RGBColor.from_string(color_rgb)
    if underline is not None:
        run.font.underline = underline
    if font_size is not None:
        run.font.size = Pt(font_size)
    else:
        run.font.size = Pt(DEFAULT_FONT_SIZE)


def set_cell_text_with_style(cell, text: str, color_rgb: Optional[str] = None, underline: Optional[bool] = None, font_size: Optional[float] = None) -> None:
    value = "" if text is None else str(text)

    if not cell.paragraphs:
        cell.text = value
        return

    clear_cell_content(cell)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(value)
    set_run_font(run, color_rgb=color_rgb, underline=underline, font_size=font_size)


def apply_font_to_paragraph(paragraph) -> None:
    for run in paragraph.runs:
        existing_size = run.font.size
        set_run_font(run)
        if existing_size is not None:
            run.font.size = existing_size


def apply_font_to_document(doc: Document) -> None:
    for style_name in ["Normal", "Hyperlink"]:
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        style.font.name = DEFAULT_FONT_NAME

    for paragraph in doc.paragraphs:
        apply_font_to_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    apply_font_to_paragraph(paragraph)

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            apply_font_to_paragraph(paragraph)
        for paragraph in section.footer.paragraphs:
            apply_font_to_paragraph(paragraph)


def set_cell_email_hyperlink(cell, email: str) -> None:
    # Capture font size from template before clearing
    font_size_pt = get_cell_font_size(cell)

    value = str(email or "").strip()
    if not value:
        set_cell_text(cell, "")
        return

    clear_cell_content(cell)
    paragraph = cell.paragraphs[0]
    part = cell.part
    relation_id = part.relate_to(f"mailto:{value}", RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)

    run_element = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")

    style = OxmlElement("w:rStyle")
    style.set(qn("w:val"), "Hyperlink")
    run_properties.append(style)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    run_properties.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(underline)

    # Apply template font size to the hyperlink run (stored as half-points in OOXML)
    if font_size_pt is not None:
        half_pts = str(int(font_size_pt * 2))
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), half_pts)
        run_properties.append(sz)
        sz_cs = OxmlElement("w:szCs")
        sz_cs.set(qn("w:val"), half_pts)
        run_properties.append(sz_cs)

    run_element.append(run_properties)

    text_element = OxmlElement("w:t")
    text_element.text = value
    run_element.append(text_element)
    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)


def _strip_paragraph_bullets(paragraph) -> None:
    """Remove any list/numbering/bullet formatting from a paragraph."""
    pPr = paragraph._p.find(qn('w:pPr'))
    if pPr is not None:
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            pPr.remove(numPr)


def clear_cell_content(cell) -> None:
    if not cell.paragraphs:
        cell.text = ""
        return

    while len(cell.paragraphs) > 1:
        paragraph = cell.paragraphs[-1]
        paragraph._element.getparent().remove(paragraph._element)

    first_paragraph = cell.paragraphs[0]
    for child in list(first_paragraph._element):
        first_paragraph._element.remove(child)


_URL_PATTERN = re.compile(r'(https?://[^\s\)]+)')


def _add_hyperlink_run(paragraph, url: str) -> None:
    """Append a clickable hyperlink run to the paragraph."""
    part = paragraph.part
    relation_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)

    run_element = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")

    style = OxmlElement("w:rStyle")
    style.set(qn("w:val"), "Hyperlink")
    run_properties.append(style)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    run_properties.append(color)

    underline_el = OxmlElement("w:u")
    underline_el.set(qn("w:val"), "single")
    run_properties.append(underline_el)

    run_element.append(run_properties)

    text_element = OxmlElement("w:t")
    text_element.text = url
    run_element.append(text_element)
    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)


def add_text_with_hyperlinks(paragraph, text: str) -> None:
    """Add text to a paragraph, converting any URLs into clickable hyperlinks."""
    pos = 0
    for match in _URL_PATTERN.finditer(text):
        before = text[pos:match.start()]
        if before:
            paragraph.add_run(before)
        _add_hyperlink_run(paragraph, match.group(1))
        pos = match.end()
    remaining = text[pos:]
    if remaining:
        paragraph.add_run(remaining)


def split_statements(text) -> list[str]:
    if not text:
        return []

    if isinstance(text, list):
        statements = []
        for item in text:
            if isinstance(item, str):
                cleaned = item.strip()
            else:
                cleaned = format_value(item).strip()
            if cleaned:
                cleaned = cleaned.lstrip("\u2022").lstrip("-").strip()
                first_line_end = cleaned.find("\n")
                if first_line_end == -1:
                    cleaned = strip_list_prefix(cleaned)
                else:
                    first_line = strip_list_prefix(cleaned[:first_line_end])
                    rest = cleaned[first_line_end:]
                    cleaned = first_line + rest
                if cleaned:
                    statements.append(cleaned)
        return statements
    elif not isinstance(text, str):
        text = format_value(text)

    statements = []
    for raw_line in str(text).replace("\r\n", "\n").splitlines():
        line = raw_line.strip().lstrip("\u2022").lstrip("-").strip()
        if not line:
            continue
        line = strip_list_prefix(line)
        if not line:
            continue
        statements.append(line)
    return statements


def split_row_items(value, preserve_list_gaps: bool = False, keep_string_block: bool = False) -> list[str]:
    if isinstance(value, list):
        statements = []
        for item in value:
            if isinstance(item, str):
                cleaned = item.strip()
            else:
                cleaned = format_value(item).strip()

            if cleaned:
                cleaned = cleaned.lstrip("\u2022").lstrip("-").strip()
                first_line_end = cleaned.find("\n")
                if first_line_end == -1:
                    cleaned = strip_list_prefix(cleaned)
                else:
                    first_line = strip_list_prefix(cleaned[:first_line_end])
                    rest = cleaned[first_line_end:]
                    cleaned = first_line + rest
                statements.append(cleaned)
            elif preserve_list_gaps:
                statements.append("")
        return statements

    if keep_string_block:
        text = format_value(value).strip()
        return [text] if text else []

    return split_statements(value)


def set_cell_lines_with_ref(cell, text, ref_tag: Optional[str] = None) -> None:
    # Capture font size from template before clearing
    font_size_pt = get_cell_font_size(cell)

    if not text:
        set_cell_text(cell, "")
        return

    if isinstance(text, list):
        text = "\n".join(format_value(item) for item in text)
    elif not isinstance(text, str):
        text = format_value(text)

    raw_lines = []
    for raw_line in str(text).replace("\r\n", "\n").splitlines():
        stripped = raw_line.strip().lstrip("\u2022").lstrip("-").strip()
        if stripped:
            raw_lines.append(stripped)

    lines = []
    parent_ends_with_colon = False
    colon_had_dotted_subs = False
    for stripped in raw_lines:
        is_sub_dotted = bool(re.match(r'^\d+\.\d+', stripped))
        is_simple_num = bool(re.match(r'^\d+[.)\-]?\s', stripped))
        is_sub = is_sub_dotted or (is_simple_num and parent_ends_with_colon and not colon_had_dotted_subs)
        cleaned = strip_list_prefix(stripped)
        if not cleaned:
            continue
        if is_sub_dotted:
            colon_had_dotted_subs = True
        lines.append((cleaned, is_sub))
        if not is_sub:
            parent_ends_with_colon = stripped.rstrip().endswith(":")
            colon_had_dotted_subs = False

    if not lines:
        set_cell_text(cell, "")
        return

    clear_cell_content(cell)
    first_paragraph = cell.paragraphs[0]
    _strip_paragraph_bullets(first_paragraph)

    for index, (statement, is_sub) in enumerate(lines):
        paragraph = first_paragraph if index == 0 else cell.add_paragraph()
        apply_bullet_formatting(paragraph, is_sub)
        add_text_with_hyperlinks(paragraph, statement)

    # Restore template font size across all new runs
    apply_font_size_to_cell(cell, font_size_pt)


def set_a2_cell_content(cell, blocks, ref_tag: str) -> None:
    # Capture font size from template before clearing
    font_size_pt = get_cell_font_size(cell)

    clear_cell_content(cell)
    first_paragraph = cell.paragraphs[0]

    for block_index, (heading, value) in enumerate(blocks):
        if block_index > 0:
            cell.add_paragraph()

        heading_paragraph = first_paragraph if block_index == 0 else cell.add_paragraph()
        heading_run = heading_paragraph.add_run(heading)
        heading_run.underline = True

        raw_lines = []
        for raw_line in str(value or "").replace("\r\n", "\n").splitlines():
            stripped = raw_line.strip().lstrip("\u2022").lstrip("-").strip()
            if stripped:
                raw_lines.append(stripped)

        has_numbered_items = any(re.match(r'^\d+', line) for line in raw_lines)

        for stripped in raw_lines:
            is_sub = bool(re.match(r'^\d+\.\d+', stripped))
            is_numbered = bool(re.match(r'^\d+', stripped))

            cleaned = strip_list_prefix(stripped)
            if not cleaned:
                continue

            value_paragraph = cell.add_paragraph()
            if is_sub or is_numbered or not has_numbered_items:
                apply_bullet_formatting(value_paragraph, is_sub)
            add_text_with_hyperlinks(value_paragraph, cleaned)
            value_paragraph.add_run(" ")
            ref_run = value_paragraph.add_run(f"({ref_tag})")
            ref_run.bold = True

    # Restore template font size across all new runs
    apply_font_size_to_cell(cell, font_size_pt)


def format_heading_text(value: str) -> str:
    text = str(value or "").replace("_", " ").strip()
    if not text:
        return ""
    return text[:1].upper() + text[1:].lower()


def strip_list_prefix(text: str) -> str:
    value = str(text or "").strip()
    if not value:
        return ""

    dotted_match = re.match(r"^(\d+(?:\.\d+)+)[.)-]?\s+(.*)$", value)
    if dotted_match:
        return dotted_match.group(2).strip()

    punctuated_match = re.match(r"^(\d+)[.)-]\s+(.*)$", value)
    if punctuated_match:
        return punctuated_match.group(2).strip()

    plain_number_match = re.match(r"^(\d+)\s+(.*)$", value)
    if plain_number_match:
        remainder = plain_number_match.group(2).strip()
        if remainder:
            first_char = remainder[0]
            if first_char.isupper():
                return remainder
    return value


def apply_bullet_formatting(paragraph, is_sub=False):
    """Applies native Word list formatting instead of manual characters."""
    try:
        paragraph.style = 'DeptBullets'
    except Exception:
        # Fallback if style is missing: manually set a hanging indent
        from docx.shared import Inches
        paragraph.paragraph_format.left_indent = Inches(0.5 if is_sub else 0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)


def set_cell_structured_needs(cell, value, ref_tag: str) -> None:
    # Capture font size from template before clearing
    font_size_pt = get_cell_font_size(cell)

    if not isinstance(value, dict):
        set_cell_bullets_with_ref(cell, format_value(value), ref_tag)
        return

    clear_cell_content(cell)
    first_paragraph = cell.paragraphs[0]
    paragraph_index = 0

    for heading, entries in value.items():
        heading_text = format_heading_text(heading)
        if not heading_text:
            continue

        heading_paragraph = first_paragraph if paragraph_index == 0 else cell.add_paragraph()
        heading_run = heading_paragraph.add_run(heading_text)
        heading_run.italic = True
        heading_run.underline = True
        paragraph_index += 1

        if not isinstance(entries, list):
            entries = [entries]

        for entry in entries:
            raw = str(entry or "").strip()
            if not raw:
                continue
            is_sub = bool(re.match(r'^\d+\.\d+', raw))
            entry_text = strip_list_prefix(raw)
            if not entry_text:
                continue
            bullet_paragraph = cell.add_paragraph()
            apply_bullet_formatting(bullet_paragraph, is_sub)
            add_text_with_hyperlinks(bullet_paragraph, entry_text)
            bullet_paragraph.add_run(" ")
            ref_run = bullet_paragraph.add_run(f"({ref_tag})")
            ref_run.bold = True
            paragraph_index += 1

    # Restore template font size across all new runs
    apply_font_size_to_cell(cell, font_size_pt)


def set_cell_bullets_with_ref(cell, text, ref_tag: Optional[str] = None) -> None:
    # Capture font size from template before clearing
    font_size_pt = get_cell_font_size(cell)

    if not text:
        set_cell_text(cell, "")
        return

    if isinstance(text, list):
        raw_items = []
        for item in text:
            s = item.strip() if isinstance(item, str) else format_value(item).strip()
            if s:
                raw_items.append(s)
    elif isinstance(text, str):
        raw_items = [line.strip() for line in text.replace("\r\n", "\n").splitlines() if line.strip()]
    else:
        raw_items = [format_value(text).strip()]

    if not raw_items:
        set_cell_text(cell, "")
        return

    items = []
    parent_ends_with_colon = False
    colon_had_dotted_subs = False

    for raw in raw_items:
        stripped = raw.lstrip("\u2022").lstrip("-").strip()
        if not stripped:
            continue

        is_sub_dotted = bool(re.match(r'^\d+\.\d+', stripped))
        is_simple_num = bool(re.match(r'^\d+[.)\-]?\s', stripped))
        is_sub = is_sub_dotted or (is_simple_num and parent_ends_with_colon and not colon_had_dotted_subs)

        cleaned = strip_list_prefix(stripped)
        if not cleaned:
            continue

        if is_sub_dotted:
            colon_had_dotted_subs = True

        items.append((cleaned, is_sub))

        if not is_sub:
            parent_ends_with_colon = stripped.rstrip().endswith(":")
            colon_had_dotted_subs = False

    if not items:
        set_cell_text(cell, "")
        return

    clear_cell_content(cell)
    first_paragraph = cell.paragraphs[0]

    for index, (statement, is_sub) in enumerate(items):
        paragraph = first_paragraph if index == 0 else cell.add_paragraph()
        apply_bullet_formatting(paragraph, is_sub)
        add_text_with_hyperlinks(paragraph, statement)

        if ref_tag:
            paragraph.add_run(" ")
            ref_run = paragraph.add_run(f"({ref_tag})")
            ref_run.bold = True

    # Restore template font size across all new runs
    apply_font_size_to_cell(cell, font_size_pt)


def set_cell_plain_lines_with_ref(cell, text, ref_tag: Optional[str] = None) -> None:
    # Capture font size from template before clearing
    font_size_pt = get_cell_font_size(cell)

    statements = split_statements(text)
    if not statements:
        set_cell_text(cell, "")
        return

    clear_cell_content(cell)
    first_paragraph = cell.paragraphs[0]
    _strip_paragraph_bullets(first_paragraph)

    for index, statement in enumerate(statements):
        paragraph = first_paragraph if index == 0 else cell.add_paragraph()
        _strip_paragraph_bullets(paragraph)
        add_text_with_hyperlinks(paragraph, statement)
        if ref_tag:
            paragraph.add_run(" ")
            ref_run = paragraph.add_run(f"({ref_tag})")
            ref_run.bold = True

    # Restore template font size across all new runs
    apply_font_size_to_cell(cell, font_size_pt)


def format_k_date(value: str) -> str:
    text = str(value or "").strip()
    if not text:
        return ""

    text = re.sub(r"\b(\d{1,2})(st|nd|rd|th)\b", r"\1", text, flags=re.IGNORECASE)
    for fmt in ("%d/%m/%Y", "%d-%m-%Y", "%d.%m.%Y", "%d/%m/%y", "%d-%m-%y", "%d.%m.%y", "%d %B %Y", "%d %b %Y", "%d %B %y", "%d %b %y"):
        try:
            return datetime.strptime(text, fmt).strftime("%d/%m/%Y")
        except ValueError:
            continue
    return text


def format_author_role_line(name: str, role: str) -> str:
    left = str(name or "").strip()
    right = str(role or "").strip()
    if left and right:
        return f"{left} - {right}"
    return left or right


def remove_table_column(table, column_index: int, skip_row_indexes=None) -> None:
    skip_row_indexes = set(skip_row_indexes or [])
    grid = getattr(table._tbl, "tblGrid", None)
    if grid is not None and len(grid.gridCol_lst) > column_index:
        grid.remove(grid.gridCol_lst[column_index])

    for row_index, row in enumerate(table.rows):
        if row_index in skip_row_indexes:
            continue
        if len(row.cells) > column_index:
            row._tr.remove(row.cells[column_index]._tc)


def fill_multiline_rows(table, start_row_index: int, column_index: int, text: str) -> None:
    if start_row_index >= len(table.rows):
        return

    parts = [line.strip() for line in str(text or "").splitlines() if line.strip()]
    template_row = table.rows[start_row_index]

    if not parts:
        if len(template_row.cells) > column_index:
            set_cell_text(template_row.cells[column_index], "")
        return

    empty_row_indexes = []
    for row_index in range(start_row_index, len(table.rows)):
        row = table.rows[row_index]
        if len(row.cells) <= column_index:
            continue
        if not str(row.cells[column_index].text or "").strip():
            empty_row_indexes.append(row_index)

    if not empty_row_indexes:
        return

    for part_index, part in enumerate(parts):
        if part_index < len(empty_row_indexes):
            target_row_index = empty_row_indexes[part_index]
        else:
            target_row_index = empty_row_indexes[-1]
            part = "\n".join([table.rows[target_row_index].cells[column_index].text.strip(), part]).strip()

        target_row = table.rows[target_row_index]
        if len(target_row.cells) > column_index:
            set_cell_text(target_row.cells[column_index], part)


def clone_blank_row_after(table, row_index: int) -> int:
    if row_index >= len(table.rows):
        return len(table.rows) - 1

    source_row = table.rows[row_index]
    new_tr = deepcopy(source_row._tr)
    source_row._tr.addnext(new_tr)
    new_row_index = row_index + 1
    new_row = table.rows[new_row_index]
    for cell in new_row.cells:
        set_cell_text(cell, "")
    return new_row_index


def fill_multiline_rows_with_ref(table, start_row_index: int, end_row_index: int, column_texts: Dict[int, Any], ref_tag: str, bullet_columns=None, ref_columns=None, plain_columns=None, preserve_list_gap_columns=None, keep_string_block_columns=None) -> int:
    """Fill table rows with multi-item data, cloning rows as needed. Returns the last row index used."""
    row_range = list(range(start_row_index, min(end_row_index, len(table.rows) - 1) + 1))
    if not row_range:
        return start_row_index

    bullet_columns = set(bullet_columns) if bullet_columns is not None else set(column_texts.keys())
    ref_columns = set(ref_columns) if ref_columns is not None else set(column_texts.keys())
    plain_columns = set(plain_columns) if plain_columns is not None else set()
    preserve_list_gap_columns = set(preserve_list_gap_columns) if preserve_list_gap_columns is not None else set()
    keep_string_block_columns = set(keep_string_block_columns) if keep_string_block_columns is not None else set()

    for row_index in row_range:
        row = table.rows[row_index]
        for column_index in column_texts.keys():
            if len(row.cells) > column_index:
                set_cell_text(row.cells[column_index], "")

    column_items = {
        column_index: split_row_items(
            text,
            preserve_list_gaps=column_index in preserve_list_gap_columns,
            keep_string_block=column_index in keep_string_block_columns,
        )
        for column_index, text in column_texts.items()
    }
    max_items = max((len(items) for items in column_items.values()), default=0)
    if max_items == 0:
        for row_index in row_range:
            row = table.rows[row_index]
            for column_index in column_texts.keys():
                if len(row.cells) > column_index:
                    set_cell_text(row.cells[column_index], "")
        return row_range[-1]

    while len(row_range) < max_items:
        row_range.append(clone_blank_row_after(table, row_range[-1]))

    usable_rows = row_range[:max_items]

    for item_index, row_index in enumerate(usable_rows):
        row = table.rows[row_index]
        for column_index, items in column_items.items():
            value = items[item_index] if item_index < len(items) else ""
            if len(row.cells) > column_index:
                column_ref = ref_tag if column_index in ref_columns else None
                if column_index in bullet_columns:
                    set_cell_bullets_with_ref(row.cells[column_index], value, column_ref)
                elif column_index in plain_columns:
                    set_cell_plain_lines_with_ref(row.cells[column_index], value, column_ref)
                else:
                    set_cell_lines_with_ref(row.cells[column_index], value, column_ref)

    for row_index in row_range[len(usable_rows):]:
        row = table.rows[row_index]
        for column_index in column_texts.keys():
            if len(row.cells) > column_index:
                set_cell_text(row.cells[column_index], "")

    return row_range[-1]


def fill_a1_section(doc: Document, education_data: Dict[str, Any], mapping_rows) -> None:
    table = find_table_by_heading(doc, "A1 SUMMARY of the child or young person's history")
    if not table or len(table.rows) < 2:
        return
    content_parts = []
    for row in mapping_rows:
        value = combine_mapping_values(education_data, row["section"], row["key"])
        if value:
            content_parts.append(value)
    set_cell_bullets_with_ref(table.rows[1].cells[0], "\n".join(content_parts), "K1")


def fill_a2_section(doc: Document, personal_data: Dict[str, Any], education_data: Dict[str, Any], mapping_rows) -> None:
    table = find_table_by_heading(doc, "A2 - Summary of the views, interests and aspirations of the child and their parent, or of the young person")
    if not table:
        return

    header_to_value_row = {
        normalize_label("Views"): 3,
        normalize_label("Strengths, Interests, passions and skills and the things that are important"): 5,
        normalize_label("Aspirations"): 7,
    }

    grouped = {}
    for row in mapping_rows:
        grouped.setdefault(row["template_column"], []).append(row)

    for template_column, rows in grouped.items():
        target_row_index = header_to_value_row.get(normalize_label(template_column))
        if target_row_index is None or target_row_index >= len(table.rows):
            continue

        blocks = []
        for row in rows:
            value = combine_mapping_values(education_data, row["section"], row["key"])
            if not value:
                continue
            label = row["template_subcolumn"] or ""
            if label:
                label = child_perspective_label(personal_data, label)
                blocks.append((label, value))
            else:
                blocks.append(("", value))

        heading_blocks = [(heading, value) for heading, value in blocks if heading]
        if heading_blocks:
            set_a2_cell_content(table.rows[target_row_index].cells[0], heading_blocks, "K1")
        else:
            set_cell_lines_with_ref(table.rows[target_row_index].cells[0], "\n\n".join(value for _, value in blocks), "K1")


def fill_primary_area_of_need_section(doc: Document, education_data: Dict[str, Any], mapping_rows, ref_tag: str = "K1") -> None:
    table = find_table_by_heading(doc, "Primary area of need")
    if not table or len(table.rows) < 2:
        return
    content_parts = []
    for row in mapping_rows:
        value = combine_mapping_values(education_data, row["section"], row["key"])
        if value:
            content_parts.append(value)
    set_cell_bullets_with_ref(table.rows[1].cells[0], "\n".join(content_parts), ref_tag)


def find_first_blank_data_row(table, start_row: int) -> int:
    for row_index in range(start_row, len(table.rows)):
        if all(not cell.text.strip() for cell in table.rows[row_index].cells):
            return row_index
    return start_row


def fill_send_section_table(table, values: Dict[str, Any]) -> None:
    if len(table.rows) < SEND_TABLE_MIN_ROWS:
        return
    set_cell_bullets_with_ref(table.rows[2].cells[0], values.get("Strengths", ""), "K1")
    set_cell_structured_needs(table.rows[4].cells[0], values.get("Needs", ""), "K1")
    if len(table.rows[5].cells) > 2:
        set_cell_bullets_with_ref(table.rows[5].cells[2], values.get("Long Term / Life Outcomes", ""), "K1")
    if len(table.rows[6].cells) > 2:
        set_cell_bullets_with_ref(table.rows[6].cells[2], values.get("Medium Term Outcomes", ""), "K1")

    fill_multiline_rows_with_ref(
        table,
        8,
        len(table.rows) - 1,
        {
            0: values.get("Special Educational Provision", ""),
            1: values.get("By Whom", ""),
            2: values.get("Frequency", ""),
            3: values.get("Which Need Does This Provision Meet", ""),
        },
        "K1",
        bullet_columns=set(),
        ref_columns={0, 1, 2, 3},
        plain_columns={0, 1, 2, 3},
        preserve_list_gap_columns={0, 1, 2, 3},
    )


def normalize_template_column(name: str) -> str:
    value = normalize_label(name or "")
    if value == normalize_label("Special Educational Provision "):
        return "special educational provision"
    if value == normalize_label("Which Need Does This Provision Meet "):
        return "which need does this provision meet"
    return value


def group_mapping_rows_by_section(mapping_rows):
    groups = {}
    for row in mapping_rows:
        section_name = str(row.get("template_section") or "").strip()
        if not section_name:
            continue
        normalized = normalize_label(section_name)
        if normalized not in groups:
            groups[normalized] = {"name": section_name, "rows": []}
        groups[normalized]["rows"].append(row)
    return list(groups.values())


def get_group_columns(rows) -> set[str]:
    return {
        normalize_template_column(row.get("template_column") or "")
        for row in rows
        if str(row.get("template_column") or "").strip()
    }


def get_group_keys(rows) -> set[str]:
    return {
        normalize_label(key_name)
        for row in rows
        for key_name in str(row.get("key") or "").split(",")
        if key_name.strip()
    }


def find_first_section_group(section_groups, required_columns=None, required_keys=None, exclude_names=None):
    required_columns = {normalize_template_column(value) for value in (required_columns or [])}
    required_keys = {normalize_label(value) for value in (required_keys or [])}
    exclude_names = {normalize_label(value) for value in (exclude_names or [])}

    for group in section_groups:
        if normalize_label(group["name"]) in exclude_names:
            continue
        group_columns = get_group_columns(group["rows"])
        group_keys = get_group_keys(group["rows"])
        if required_columns and not required_columns.issubset(group_columns):
            continue
        if required_keys and not required_keys.issubset(group_keys):
            continue
        return group
    return None


def find_all_section_groups(section_groups, required_columns=None, required_keys=None, exclude_names=None):
    required_columns = {normalize_template_column(value) for value in (required_columns or [])}
    required_keys = {normalize_label(value) for value in (required_keys or [])}
    exclude_names = {normalize_label(value) for value in (exclude_names or [])}

    matches = []
    for group in section_groups:
        if normalize_label(group["name"]) in exclude_names:
            continue
        group_columns = get_group_columns(group["rows"])
        group_keys = get_group_keys(group["rows"])
        if required_columns and not required_columns.issubset(group_columns):
            continue
        if required_keys and not required_keys.issubset(group_keys):
            continue
        matches.append(group)
    return matches


def build_mapping_values(data: Dict[str, Any], rows, outcome_fallbacks=None, column_aliases=None) -> Dict[str, Any]:
    outcome_fallbacks = outcome_fallbacks or {}
    column_aliases = {normalize_template_column(key): value for key, value in (column_aliases or {}).items()}
    values = {}

    for row in rows:
        template_column_normalized = normalize_template_column(row.get("template_column") or "")
        key = (row.get("template_subcolumn") or row.get("template_column") or "").strip()
        if not key:
            key = outcome_fallbacks.get(normalize_label(row.get("key") or ""), "")
        alias_key = column_aliases.get(template_column_normalized)
        if alias_key:
            key = alias_key
        if not key:
            continue
        values[key] = resolve_mapping_value(data, row.get("section") or "", row.get("key") or "")

    return values


def fill_k_table_entry(doc: Document, mapping_rows, data: Dict[str, Any]) -> None:
    if not mapping_rows:
        return

    table = find_table_by_heading(doc, K_TABLE_HEADING)
    if not table or len(table.rows) < 2:
        return

    row_label = str(mapping_rows[0].get("template_column") or "").strip()
    if not row_label:
        return

    row_index = find_k_row_index(table, row_label)
    if row_index is None:
        return

    date_value = ""
    author_parts = []

    for mapping_row in mapping_rows:
        value = resolve_mapping_value(data, mapping_row.get("section") or "", mapping_row.get("key") or "")
        text = format_value(value).strip()
        detail_label = normalize_label(mapping_row.get("template_subcolumn") or "")
        key_label = normalize_label(mapping_row.get("key") or "")

        if detail_label == normalize_label("Date of Document") or key_label in {normalize_label("Date"), normalize_label("Advice_Date")}:
            if text:
                date_value = text
            continue

        if detail_label == normalize_label("Author & Role"):
            if text:
                author_parts.append(text)

    if not author_parts:
        fallback_name = str(data.get("Name") or "").strip() if isinstance(data, dict) else ""
        fallback_role = str(data.get("Role") or "").strip() if isinstance(data, dict) else ""
        author_parts = [part for part in [fallback_name, fallback_role] if part]

    target_row = table.rows[row_index]
    if len(target_row.cells) > 2:
        set_cell_text(target_row.cells[2], format_k_date(date_value))
    if len(target_row.cells) > 3:
        author_name = author_parts[0] if len(author_parts) > 0 else ""
        author_role = author_parts[1] if len(author_parts) > 1 else ""
        if len(author_parts) > 2 and not author_role:
            author_role = " ".join(author_parts[1:])
        set_cell_text(target_row.cells[3], format_author_role_line(author_name, author_role))


def fill_send_sections(doc: Document, education_data: Dict[str, Any], mapping_rows) -> None:
    section_groups = {}
    for row in mapping_rows:
        section_groups.setdefault(row["template_section"], []).append(row)

    for template_section, rows in section_groups.items():
        table = find_table_by_heading(doc, template_section, min_rows=SEND_TABLE_MIN_ROWS)
        if not table:
            continue
        values = {}
        for row in rows:
            template_column = normalize_template_column(row["template_column"])
            key = row["template_subcolumn"] or row["template_column"] or ""
            key = key.strip()
            if normalize_template_column(row["template_column"]) == "outcomes":
                key = row["template_subcolumn"]
                if not key:
                    if row["key"] == "Long_Term_Outcomes":
                        key = "Long Term / Life Outcomes"
                    elif row["key"] == "Medium_Term_Outcomes":
                        key = "Medium Term Outcomes"
            elif normalize_template_column(row["template_column"]) == "special educational provision":
                key = "Special Educational Provision"
            elif normalize_template_column(row["template_column"]) == "which need does this provision meet":
                key = "Which Need Does This Provision Meet"
            value = resolve_mapping_value(education_data, row["section"], row["key"])
            values[key] = value
        fill_send_section_table(table, values)


def fill_k_table(doc: Document, education_data: Dict[str, Any]) -> None:
    fill_k_table_entry(doc, [], education_data)


def fill_k_row(table, row_index: int, advice_data: Dict[str, Any]) -> None:
    if not isinstance(advice_data, dict) or row_index >= len(table.rows):
        return
    row = table.rows[row_index]
    if len(row.cells) > 2:
        set_cell_text(row.cells[2], format_k_date(str(advice_data.get("Date") or "")))
    if len(row.cells) > 3:
        author_name = str(advice_data.get("Name") or "").strip()
        author_role = str(advice_data.get("Role") or "").strip()
        set_cell_text(row.cells[3], format_author_role_line(author_name, author_role))


def find_k_row_index(table, label: str) -> Optional[int]:
    target = normalize_label(label)
    for row_index, row in enumerate(table.rows):
        if any(target == normalize_label(cell.text) for cell in row.cells if cell.text.strip()):
            return row_index
    return None


def fill_socialcare_sections(docx_path: str, docx_out: str, socialcare_data: Dict[str, Any], mapping_path: str) -> None:
    mapping_rows = load_education_mapping(mapping_path, SOCIAL_CARE_DETAILS_SHEET)
    doc = Document(docx_path)
    section_groups = group_mapping_rows_by_section(mapping_rows)

    root_group = find_first_section_group(section_groups, required_columns={"Strengths", "Needs", "Outcomes"})
    provision_groups = find_all_section_groups(
        section_groups,
        required_columns={"Social Care Provision", "By Whom", "Frequency", "Which Need Does This Provision Meet"},
    )
    k_group = find_first_section_group(section_groups, required_keys={"date"}, exclude_names={root_group["name"]} if root_group else None)

    social_table_heading = root_group["name"] if root_group else ""
    social_table = find_table_by_heading(doc, social_table_heading)
    if social_table and root_group:
        root_values = build_mapping_values(socialcare_data, root_group["rows"])

        set_cell_bullets_with_ref(social_table.rows[2].cells[0], root_values.get("Strengths", ""), "K3")
        set_cell_bullets_with_ref(social_table.rows[4].cells[0], root_values.get("Needs", ""), "K3")
        set_cell_bullets_with_ref(social_table.rows[5].cells[2], root_values.get("Long Term / Life Outcomes", ""), "K3")
        set_cell_bullets_with_ref(social_table.rows[6].cells[2], root_values.get("Medium Term Outcomes", ""), "K3")

        H1_START = 9
        H1_DEFAULT_END = 14
        H1_TO_H2_GAP = 3
        H2_DEFAULT_SPAN = 5

        h1_last_row = H1_DEFAULT_END

        if len(provision_groups) > 0:
            h1_values = build_mapping_values(
                socialcare_data,
                provision_groups[0]["rows"],
                column_aliases={
                    "Social Care Provision": "Social Care Provision",
                    "Which Need Does This Provision Meet": "Which Need Does This Provision Meet",
                },
            )
            h1_last_row = fill_multiline_rows_with_ref(
                social_table,
                H1_START,
                H1_DEFAULT_END,
                {
                    0: h1_values.get("Social Care Provision", ""),
                    1: h1_values.get("By Whom", ""),
                    2: h1_values.get("Frequency", ""),
                    3: h1_values.get("Which Need Does This Provision Meet", ""),
                },
                "K3",
                bullet_columns=set(),
                ref_columns={0, 1, 2, 3},
                plain_columns={0, 1, 2, 3},
                keep_string_block_columns={0, 1, 2, 3},
                preserve_list_gap_columns={0, 1, 2, 3},
            )

        if len(provision_groups) > 1:
            h2_values = build_mapping_values(
                socialcare_data,
                provision_groups[1]["rows"],
                column_aliases={
                    "Social Care Provision": "Social Care Provision",
                    "Which Need Does This Provision Meet": "Which Need Does This Provision Meet",
                },
            )
            h1_extra_rows = max(0, h1_last_row - H1_DEFAULT_END)
            h2_start = H1_DEFAULT_END + H1_TO_H2_GAP + h1_extra_rows
            h2_end = h2_start + H2_DEFAULT_SPAN
            fill_multiline_rows_with_ref(
                social_table,
                h2_start,
                h2_end,
                {
                    0: h2_values.get("Social Care Provision", ""),
                    1: h2_values.get("By Whom", ""),
                    2: h2_values.get("Frequency", ""),
                    3: h2_values.get("Which Need Does This Provision Meet", ""),
                },
                "K3",
                bullet_columns=set(),
                ref_columns={0, 1, 2, 3},
                plain_columns={0, 1, 2, 3},
                keep_string_block_columns={0, 1, 2, 3},
                preserve_list_gap_columns={0, 1, 2, 3},
            )

    if k_group:
        fill_k_table_entry(doc, k_group["rows"], socialcare_data)

    apply_font_to_document(doc)
    doc.save(docx_out)


def fill_health_sections(docx_path: str, docx_out: str, health_data: Dict[str, Any], mapping_path: str) -> None:
    mapping_rows = load_education_mapping(mapping_path, HEALTH_DETAILS_SHEET)
    doc = Document(docx_path)
    section_groups = group_mapping_rows_by_section(mapping_rows)

    health_group = find_first_section_group(section_groups, required_columns={"Needs", "Outcomes", "Health Provision"})
    k_group = find_first_section_group(section_groups, required_keys={"advice date"}, exclude_names={health_group["name"]} if health_group else None)

    health_table = find_table_by_heading(doc, health_group["name"] if health_group else "")
    if health_table and health_group:
        values = build_mapping_values(
            health_data,
            health_group["rows"],
            outcome_fallbacks={normalize_label("Outcomes"): "Medium Term Outcomes"},
            column_aliases={"Health Provision": "Health Provision"},
        )

        set_cell_bullets_with_ref(health_table.rows[2].cells[0], values.get("Needs", ""), "K2")
        # Health mapping currently requires only Medium Term Outcomes.
        set_cell_bullets_with_ref(health_table.rows[4].cells[2], values.get("Medium Term Outcomes", ""), "K2")
        fill_multiline_rows_with_ref(
            health_table,
            6,
            11,
            {0: values.get("Health Provision", "")},
            "K2",
            bullet_columns=set(),
            ref_columns={0},
            plain_columns={0},
        )

    if k_group:
        fill_k_table_entry(doc, k_group["rows"], health_data)

    apply_font_to_document(doc)
    doc.save(docx_out)


def build_default_output_path(docx_path: str, personal_data: Dict[str, Any]) -> str:
    full_name = value_for(personal_data, "name").strip() or value_for(personal_data, "preferred_name").strip()
    safe_name = re.sub(r'[<>:"/\\|?*]+', " ", full_name)
    safe_name = re.sub(r"\s+", " ", safe_name).strip() or "OUTPUT"
    output_dir = os.path.join(os.path.dirname(docx_path), "output")
    os.makedirs(output_dir, exist_ok=True)
    return os.path.join(output_dir, f"{safe_name} Draft EHCP.docx")


def resolve_output_docx_path(docx_path: str, personal_data: Dict[str, Any], requested_output: Optional[str]) -> str:
    if not requested_output:
        return build_default_output_path(docx_path, personal_data)

    if os.path.isabs(requested_output):
        output_dir = os.path.dirname(requested_output)
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)
        return requested_output

    output_dir = os.path.join(os.path.dirname(docx_path), "output")
    os.makedirs(output_dir, exist_ok=True)
    return os.path.join(output_dir, requested_output)


def resolve_input_json_paths(args) -> Dict[str, str]:
    if not args.input_folder:
        return {
            "json": args.json,
            "education_json": args.education_json,
            "socialcare_json": args.socialcare_json,
            "health_json": args.health_json,
        }

    return {
        "json": os.path.join(args.input_folder, "personal_details.json"),
        "education_json": os.path.join(args.input_folder, "education_advice.json"),
        "socialcare_json": os.path.join(args.input_folder, "socialcare_advice.json"),
        "health_json": os.path.join(args.input_folder, "health_advice.json"),
    }


def fill_education_sections(docx_path: str, docx_out: str, personal_data: Dict[str, Any], education_data: Dict[str, Any], mapping_path: str) -> None:
    mapping_rows = load_education_mapping(mapping_path)
    doc = Document(docx_path)
    section_groups = group_mapping_rows_by_section(mapping_rows)
    a1_group = find_first_section_group(section_groups, required_keys={"Journey_so_far"})
    a2_group = find_first_section_group(section_groups, required_columns={"Views", "Aspirations"})
    primary_area_of_need_group = find_first_section_group(section_groups, required_keys={"Primary_area_of_need"})
    k_group = find_first_section_group(section_groups, required_keys={"date", "name", "role"})
    excluded_names = {
        group["name"]
        for group in [a1_group, a2_group, k_group]
        if group
    }
    send_rows = [
        row
        for group in section_groups
        if group["name"] not in excluded_names
        for row in group["rows"]
    ]

    fill_a1_section(doc, education_data, a1_group["rows"] if a1_group else [])
    fill_a2_section(doc, personal_data, education_data, a2_group["rows"] if a2_group else [])
    fill_primary_area_of_need_section(doc, education_data, primary_area_of_need_group["rows"] if primary_area_of_need_group else [])
    fill_send_sections(doc, education_data, send_rows)
    if k_group:
        fill_k_table_entry(doc, k_group["rows"], education_data)

    apply_font_to_document(doc)
    doc.save(docx_out)


def fill_template_streamlit(
    template_docx: str,
    output_docx: Optional[str],
    personal_json_path: str,
    education_json_path: str,
    health_json_path: str,
    socialcare_json_path: str,
    mapping_workbook: str,
) -> str:
    personal_data = load_optional_json(personal_json_path)
    personal_mapping = load_personal_mapping(mapping_workbook)
    resolved_output_docx = resolve_output_docx_path(template_docx, personal_data, output_docx)

    create_filled_template_docx(template_docx, resolved_output_docx, personal_data, personal_mapping)

    if os.path.exists(education_json_path):
        education_data = load_optional_json(education_json_path)
        fill_education_sections(resolved_output_docx, resolved_output_docx, personal_data, education_data, mapping_workbook)

    if os.path.exists(socialcare_json_path):
        socialcare_data = load_optional_json(socialcare_json_path)
        fill_socialcare_sections(resolved_output_docx, resolved_output_docx, socialcare_data, mapping_workbook)

    if os.path.exists(health_json_path):
        health_data = load_optional_json(health_json_path)
        fill_health_sections(resolved_output_docx, resolved_output_docx, health_data, mapping_workbook)

    return resolved_output_docx


def main():
    start_time = time.perf_counter()
    parser = argparse.ArgumentParser(description="Convert DOCX -> MD, fill MD using JSON, and write filled DOCX.")
    parser.add_argument("--docx", required=True, help="Path to source DOCX file")
    parser.add_argument("--input-folder", default=None, help="Path to a child input folder containing personal_details.json, education_advice.json, socialcare_advice.json, and health_advice.json")
    parser.add_argument("--json", default=os.path.join("input", "personal_details.json"), help="Path to personal details JSON file")
    parser.add_argument("--education-json", default=os.path.join("input", "education_advice.json"), help="Path to extracted education JSON file")
    parser.add_argument("--socialcare-json", default=os.path.join("input", "socialcare_advice.json"), help="Path to social care advice JSON file")
    parser.add_argument("--health-json", default=os.path.join("input", "health_advice.json"), help="Path to health advice JSON file")
    parser.add_argument("--mapping-workbook", default="ehcp_mapping.xlsx", help="Path to workbook with mapping sheets for personal, education, health, and social care")
    parser.add_argument("--education-mapping", default=None, help="Deprecated alias for mapping workbook path")
    parser.add_argument("--md", default="output.md", help="Path to save converted markdown")
    parser.add_argument("--meta", default="metadata.json", help="Path to save metadata about the docx")
    parser.add_argument("--filled-md", default="filled_output.md", help="Path to save filled markdown")
    parser.add_argument("--out-docx", default=None, help="Path to save filled docx")

    args = parser.parse_args()

    json_paths = resolve_input_json_paths(args)

    if not os.path.exists(args.docx):
        raise FileNotFoundError(f"DOCX file not found: {args.docx}")
    if not os.path.exists(json_paths["json"]):
        raise FileNotFoundError(f"JSON file not found: {json_paths['json']}")

    print("Converting DOCX to Markdown...")
    docx_to_markdown(args.docx, args.md, args.meta)

    print("Loading JSON replacements...")
    data = load_json(json_paths["json"])
    mapping_workbook = args.mapping_workbook
    if not os.path.exists(mapping_workbook) and args.education_mapping:
        mapping_workbook = args.education_mapping

    personal_mapping = load_personal_mapping(mapping_workbook)
    education_data = None
    socialcare_data = None
    health_data = None
    if os.path.exists(json_paths["education_json"]):
        print("Loading education JSON...")
        education_data = load_optional_json(json_paths["education_json"])
    if os.path.exists(json_paths["socialcare_json"]):
        print("Loading social care JSON...")
        socialcare_data = load_optional_json(json_paths["socialcare_json"])
    if os.path.exists(json_paths["health_json"]):
        print("Loading health JSON...")
        health_data = load_optional_json(json_paths["health_json"])

    output_docx = resolve_output_docx_path(args.docx, data, args.out_docx)

    print("Filling Markdown with JSON values (best-effort replacements)...")
    fill_markdown(args.md, args.filled_md, data)

    print("Creating structured filled markdown according to mapping...")
    create_structured_markdown(args.filled_md, data, personal_mapping)

    print("Filling the original DOCX template while preserving its layout...")
    create_filled_template_docx(args.docx, output_docx, data, personal_mapping)

    if education_data is not None and os.path.exists(mapping_workbook):
        print("Filling education sections using the education mapping...")
        fill_education_sections(output_docx, output_docx, data, education_data, mapping_workbook)

    if socialcare_data is not None and os.path.exists(mapping_workbook):
        print("Filling social care sections using the social care mapping...")
        fill_socialcare_sections(output_docx, output_docx, socialcare_data, mapping_workbook)

    if health_data is not None and os.path.exists(mapping_workbook):
        print("Filling health sections using the health mapping...")
        fill_health_sections(output_docx, output_docx, health_data, mapping_workbook)

    elapsed_seconds = time.perf_counter() - start_time
    print(f"Done. Time taken: {elapsed_seconds:.2f} seconds")


if __name__ == "__main__":
    main()