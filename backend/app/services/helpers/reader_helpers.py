import json
import re
import os
import asyncio
import time
import contextvars
from functools import lru_cache
from dotenv import load_dotenv

from azure.ai.documentintelligence import DocumentIntelligenceClient
from azure.core.credentials import AzureKeyCredential

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import fitz  # PyMuPDF for PDF text extraction

from openai import AsyncAzureOpenAI
from azure.identity import get_bearer_token_provider

from app.services.blob_storage import upload_json_to_blob, is_blob_storage_enabled
from app.settings import (
    USE_MANAGED_IDENTITY,
    get_azure_credential,
)


# =========================================================
# Load Environment Variables
# =========================================================

load_dotenv()

endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
api_key = os.getenv("AZURE_OPENAI_API_KEY")
deployment = os.getenv("AZURE_OPENAI_DEPLOYMENT")
api_version = os.getenv("AZURE_OPENAI_API_VERSION")

di_endpoint = os.getenv("AZURE_DI_ENDPOINT") or os.getenv("AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT")
di_key = os.getenv("AZURE_DI_KEY") or os.getenv("AZURE_DOCUMENT_INTELLIGENCE_KEY")

DI_TIMEOUT_SECONDS = int(os.getenv("DI_TIMEOUT_SECONDS", "180"))
LLM_TIMEOUT_SECONDS = int(os.getenv("LLM_TIMEOUT_SECONDS", "240"))
DI_CONCURRENCY = int(os.getenv("DI_CONCURRENCY", "2"))

SUPPORTED_EXTENSIONS = [".docx", ".pdf"]


# =========================================================
# Token Usage Accumulator (per-request via contextvars)
# =========================================================

_token_accumulator: contextvars.ContextVar[dict | None] = contextvars.ContextVar(
    "_token_accumulator", default=None
)


def start_token_tracking() -> dict:
    """Start tracking token usage for the current async context.
    Returns the accumulator dict that will be updated in-place."""
    acc = {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0}
    _token_accumulator.set(acc)
    return acc


def get_token_usage() -> dict:
    """Get accumulated token usage for the current context."""
    acc = _token_accumulator.get()
    return acc if acc else {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0}


def _accumulate_tokens(usage: dict):
    """Add token usage to the current context accumulator (if active)."""
    acc = _token_accumulator.get()
    if acc and usage:
        acc["prompt_tokens"] += usage.get("prompt_tokens", 0)
        acc["completion_tokens"] += usage.get("completion_tokens", 0)
        acc["total_tokens"] += usage.get("total_tokens", 0)


# =========================================================
# Shared Helpers
# =========================================================

@lru_cache(maxsize=1)
def get_document_intelligence_client() -> DocumentIntelligenceClient:
    if USE_MANAGED_IDENTITY:
        return DocumentIntelligenceClient(
            endpoint=di_endpoint,
            credential=get_azure_credential(),
        )
    return DocumentIntelligenceClient(
        endpoint=di_endpoint,
        credential=AzureKeyCredential(di_key),
    )


def create_openai_client() -> AsyncAzureOpenAI:
    """Create an AsyncAzureOpenAI client for LLM calls."""
    if USE_MANAGED_IDENTITY:
        credential = get_azure_credential()
        token_provider = get_bearer_token_provider(
            credential, "https://cognitiveservices.azure.com/.default"
        )
        return AsyncAzureOpenAI(
            azure_endpoint=endpoint,
            azure_ad_token_provider=token_provider,
            api_version=api_version,
        )
    return AsyncAzureOpenAI(
        azure_endpoint=endpoint,
        api_key=api_key,
        api_version=api_version,
    )


@lru_cache(maxsize=8)
def load_text_file(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()


@lru_cache(maxsize=8)
def load_schema_json(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8") as f:
        return json.dumps(json.load(f), indent=2)


def safe_parse_json(json_data: str) -> dict:
    """Parse JSON from agent output, handling control chars and markdown fences."""
    cleaned = json_data.strip()

    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
        cleaned = re.sub(r"\s*```$", "", cleaned)
        cleaned = cleaned.strip()

    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        pass

    cleaned = cleaned.replace("\r\n", "\\n").replace("\r", "\\n")
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        pass

    cleaned = cleaned.replace("\n", "\\n").replace("\t", "\\t")
    return json.loads(cleaned)


# Provision-table fields that should have clean plain text (no bullets/numbering)
_PROVISION_FIELDS = {"Provision", "By_Whom", "Frequency", "Met_Need"}

# Patterns to strip at the start of lines inside provision field values
_BULLET_PREFIX_RE = re.compile(
    r"^[\s]*(?:[•·–—\-]\s*|\d+[\.\)]\s*|\d+\.\d+[\.\)]*\s*)",
    re.MULTILINE,
)

# Pattern to find URLs that may have been broken by DI with spaces/newlines
_BROKEN_URL_RE = re.compile(
    r'(https?://\S+?)(\s+)(/\S+)',
)


def _repair_broken_urls(text: str) -> str:
    """Fix URLs broken by Document Intelligence inserting spaces or line breaks,
    and prefix bare www. URLs with https://."""
    # Repeatedly fix space-broken URL segments (e.g. "documents /education" -> "documents/education")
    result = text
    while True:
        fixed = _BROKEN_URL_RE.sub(r'\1\3', result)
        if fixed == result:
            break
        result = fixed
    # Also fix URLs split across lines: http://...\n...continuation
    result = re.sub(r'(https?://[^\s\)]+)\n([^\s\)]+)', r'\1\2', result)
    # Prefix bare www. URLs with https://
    result = re.sub(r'(?<![/\w])www\.', 'https://www.', result)
    # Strip trailing sentence punctuation (. , ; :) that is not part of the URL
    result = re.sub(r'(https?://\S+?)[.,;:]+(?=\s|$)', r'\1', result)
    return result


def _clean_provision_value(value: str) -> str:
    """Remove bullet/numbering prefixes from lines within a provision field value."""
    return _BULLET_PREFIX_RE.sub("", value).strip()


# Regex to detect numbered bullet prefixes like "1 ", "2.1 ", "3.1.2 "
_NUMBERED_PREFIX_RE = re.compile(r"^\d+(\.\d+)*\s")

# Sections whose array fields contain narrative paragraphs (not bullets) and
# should be merged into single elements when the LLM splits them by sentence.
_PARAGRAPH_SECTIONS = {"Pupil_Views", "Parent_Carer_Views"}

# Fields within those sections that contain narrative text (not bullet lists)
_PARAGRAPH_FIELDS = {
    "How_I_communicate",
    "Strengths_Interests_Passions_Skills",
    "Future_Aspirations",
    "Other_things_Id_like_people_to_know",
    "Family_Future_Aspirations",
    "Other_things_parents_would_like_people_to_know",
    "Community_support",
}


def _merge_paragraph_arrays(data: dict) -> dict:
    """Merge arrays that were incorrectly split into per-sentence elements back into
    single paragraph elements.  Only applies to narrative fields in Views sections.
    If any element has a numbered prefix (e.g. '1 ...'), the array is left as-is
    because it's a genuine bullet list."""
    for section_key in _PARAGRAPH_SECTIONS:
        section = data.get(section_key)
        if not isinstance(section, dict):
            continue
        for field_key in _PARAGRAPH_FIELDS:
            val = section.get(field_key)
            if not isinstance(val, list) or len(val) <= 1:
                continue
            # If any element looks like a numbered bullet, leave it alone
            if any(_NUMBERED_PREFIX_RE.match(str(item)) for item in val if isinstance(item, str)):
                continue
            # All elements are plain paragraph text — join into one element
            merged = " ".join(str(item).strip() for item in val if isinstance(item, str) and item.strip())
            if merged:
                section[field_key] = [merged]
    return data


def post_process_json(data: dict) -> dict:
    """Clean up extracted JSON: merge split paragraphs, strip bullets from provision table fields, fix broken URLs."""

    def _fix_urls_in_value(val):
        """Recursively fix broken URLs in any string value."""
        if isinstance(val, str):
            return _repair_broken_urls(val)
        if isinstance(val, list):
            return [_fix_urls_in_value(item) for item in val]
        if isinstance(val, dict):
            return {k: _fix_urls_in_value(v) for k, v in val.items()}
        return val

    # Fix broken URLs throughout the entire output
    data = _fix_urls_in_value(data)

    # Merge paragraph arrays that LLM split into per-sentence elements
    data = _merge_paragraph_arrays(data)

    for section_key, section_val in data.items():
        if not isinstance(section_val, dict):
            continue
        for field_key, field_val in section_val.items():
            if field_key not in _PROVISION_FIELDS:
                continue
            if isinstance(field_val, list):
                section_val[field_key] = [
                    _clean_provision_value(item) if isinstance(item, str) else item
                    for item in field_val
                ]
            elif isinstance(field_val, str):
                section_val[field_key] = _clean_provision_value(field_val)

    # Enforce equal array lengths for provision table sections
    data = _enforce_provision_array_lengths(data)

    return data


def _enforce_provision_array_lengths(data: dict) -> dict:
    """Ensure Provision, By_Whom, Frequency, Met_Need arrays have equal lengths.

    If Provision is longer than the other arrays, it likely contains values
    that leaked from Met_Need or other columns.  Truncate Provision to match.
    If any shorter array needs padding, pad with null.
    """
    provision_sections = ["H1_Social_Care_Provision", "H2_Social_Care_Provision"]
    array_fields = ["Provision", "By_Whom", "Frequency", "Met_Need"]

    for section_key in provision_sections:
        section = data.get(section_key)
        if not isinstance(section, dict):
            continue

        lengths = {}
        for field in array_fields:
            val = section.get(field)
            if isinstance(val, list):
                lengths[field] = len(val)

        if not lengths:
            continue

        all_lens = list(lengths.values())
        if len(set(all_lens)) <= 1:
            continue  # already aligned

        # Use the most common length (mode) among the non-Provision arrays as
        # the target — By_Whom, Frequency, Met_Need are less likely to have
        # extra entries since their values are short/structured.
        non_provision_lens = [lengths[f] for f in ("By_Whom", "Frequency", "Met_Need") if f in lengths]
        if non_provision_lens:
            from collections import Counter
            target_len = Counter(non_provision_lens).most_common(1)[0][0]
        else:
            target_len = min(all_lens)

        for field in array_fields:
            val = section.get(field)
            if not isinstance(val, list):
                continue
            if len(val) > target_len:
                print(f"  [PostProcess] {section_key}.{field}: truncating {len(val)} → {target_len} entries")
                section[field] = val[:target_len]
            elif len(val) < target_len:
                print(f"  [PostProcess] {section_key}.{field}: padding {len(val)} → {target_len} entries")
                section[field] = val + [None] * (target_len - len(val))

    return data


# =========================================================
# Document Intelligence: Analyze
# =========================================================

def analyze_with_doc_intelligence(file_path: str) -> str:
    """Analyze a local file using Azure Document Intelligence layout model."""
    client = get_document_intelligence_client()

    with open(file_path, "rb") as f:
        poller = client.begin_analyze_document(
            model_id="prebuilt-layout",
            body=f,
            content_type="application/octet-stream",
        )

    result = poller.result()
    document_text = result.content

    # Supplement with table cells that DI parsed but didn't include in content
    if hasattr(result, 'tables') and result.tables:
        di_table_lines = []
        for table in result.tables:
            for cell in table.cells:
                cell_text = cell.content.strip() if cell.content else ""
                if cell_text and len(cell_text) > 5 and cell_text not in document_text:
                    di_table_lines.append(cell_text)
        if di_table_lines:
            document_text += "\n" + "\n".join(di_table_lines)
            print(f"  Appended {len(di_table_lines)} DI table cell(s) not in main content.")

    return document_text


# =========================================================
# Local DOCX Fallback (when DI fails)
# =========================================================

def extract_docx_text_locally(docx_path: str) -> str:
    """Extract text from DOCX using python-docx (fallback when DI is unavailable)."""
    doc = Document(docx_path)
    parts = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cell_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cell_texts:
                parts.append(" | ".join(cell_texts))

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)
        for paragraph in section.footer.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)

    return "\n".join(parts)


def extract_pdf_text_locally(pdf_path: str) -> str:
    """Extract text from PDF using PyMuPDF (fallback when DI is unavailable)."""
    doc = fitz.open(pdf_path)
    parts = []
    for page in doc:
        text = page.get_text().strip()
        if text:
            parts.append(text)
    doc.close()
    return "\n".join(parts)


def extract_text_locally(file_path: str) -> str:
    """Extract text locally from a DOCX or PDF file (fallback when DI is unavailable)."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_pdf_text_locally(file_path)
    elif ext == ".docx":
        return extract_docx_text_locally(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}")


# =========================================================
# Hyperlink Extraction
# =========================================================

def extract_hyperlinks(docx_path: str) -> list:
    """Extract hyperlinks from a DOCX file, returning a list of (display_text, url) tuples."""
    doc = Document(docx_path)
    seen = set()
    unique = []

    def _scan_paragraphs(paragraphs):
        for paragraph in paragraphs:
            for child in paragraph._element:
                if child.tag.endswith("}hyperlink"):
                    r_id = child.get(
                        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"
                    )
                    if r_id and r_id in paragraph.part.rels:
                        rel = paragraph.part.rels[r_id]
                        if rel.reltype == RT.HYPERLINK:
                            url = rel._target
                            # Only get text from w:t elements to avoid duplication
                            ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
                            display = "".join(
                                t.text for t in child.iter(f"{ns}t")
                                if t.text
                            )
                            if url and url.startswith("http") and url not in seen:
                                seen.add(url)
                                unique.append((display.strip(), url))

    _scan_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _scan_paragraphs(cell.paragraphs)

    return unique


def append_hyperlink_section(document_text: str, docx_path: str) -> str:
    """Convert DOCX hyperlink display text into markdown links and fix bare www URLs."""
    if not docx_path.lower().endswith(".docx"):
        return document_text

    hyperlinks = extract_hyperlinks(docx_path)

    # Convert hyperlink display text to markdown link format [display text](url)
    for display_text, url in hyperlinks:
        if not display_text or not url:
            continue
        # Skip if display text IS itself a URL (already usable)
        if display_text.startswith("http") or display_text.startswith("www."):
            continue
        # Replace display text with markdown link
        md_link = f"[{display_text}]({url})"
        if display_text in document_text and md_link not in document_text:
            document_text = document_text.replace(display_text, md_link, 1)

    # Prefix bare www. URLs with https:// to make them proper links
    document_text = re.sub(
        r'(?<![/\w])www\.',
        'https://www.',
        document_text,
    )

    return document_text


# =========================================================
# Patch paragraph breaks inside DOCX table cells
# =========================================================

def patch_table_paragraph_breaks(document_text: str, docx_path: str) -> str:
    """Re-insert paragraph breaks that DI flattened inside table cells.

    Document Intelligence joins paragraphs within a single table cell with a
    single newline.  This function reads the DOCX via python-docx, identifies
    cells whose paragraphs were merged, and replaces the single-newline joins
    with double-newline separators so the LLM can distinguish paragraphs.
    """
    if not docx_path.lower().endswith(".docx"):
        return document_text

    doc = Document(docx_path)
    patches_applied = 0

    def _process_table(table):
        nonlocal document_text, patches_applied
        for row in table.rows:
            for cell in row.cells:
                paras = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
                if len(paras) < 2:
                    continue
                flat = "\n".join(paras)
                spaced = "\n\n".join(paras)
                # Only patch substantial blocks to avoid false matches
                if len(flat) > 50 and flat in document_text and spaced not in document_text:
                    document_text = document_text.replace(flat, spaced, 1)
                    patches_applied += 1
                # Handle nested tables
                for nested in cell.tables:
                    _process_table(nested)

    for table in doc.tables:
        _process_table(table)

    if patches_applied:
        print(f"  [Paragraph-patch] Restored paragraph breaks in {patches_applied} table cell(s).")
    return document_text


# =========================================================
# Patch paragraph breaks in PDF text blocks
# =========================================================

def patch_pdf_paragraph_breaks(document_text: str, pdf_path: str) -> str:
    """Re-insert paragraph breaks that DI flattened inside PDF text blocks.

    DI merges paragraphs within a single text block, removing whitespace-only
    separator lines.  This function uses PyMuPDF to detect blocks that contain
    blank / whitespace-only lines (paragraph separators) and splits them into
    groups of real text lines.  The last line of one paragraph and the first
    line of the next paragraph are then used as anchors to insert ``\\n\\n``
    into the DI text.
    """
    if not pdf_path.lower().endswith(".pdf"):
        return document_text

    # DI normalises smart quotes to ASCII; PyMuPDF preserves them.
    _SMART_QUOTES = str.maketrans({
        "\u2018": "'", "\u2019": "'",   # single curly quotes
        "\u201C": '"', "\u201D": '"',   # double curly quotes
    })

    def _normalise(text: str) -> str:
        return text.translate(_SMART_QUOTES)

    doc = fitz.open(pdf_path)
    patches_applied = 0

    for page in doc:
        blocks = page.get_text("dict")["blocks"]
        for block in blocks:
            if block.get("type") != 0:
                continue
            lines = block.get("lines", [])
            if len(lines) < 3:
                continue

            # Classify each line as blank or content
            line_texts = []
            for line in lines:
                text = "".join(s["text"] for s in line["spans"])
                line_texts.append(text)

            # Collect paragraph groups (split on blank lines)
            groups = []
            current_group = []
            for lt in line_texts:
                if lt.strip() == "":
                    if current_group:
                        groups.append(current_group)
                        current_group = []
                else:
                    current_group.append(lt.rstrip())
            if current_group:
                groups.append(current_group)

            if len(groups) < 2:
                continue

            # Build the tail of one paragraph and head of the next
            for i in range(len(groups) - 1):
                tail = _normalise(groups[i][-1].rstrip())
                head = _normalise(groups[i + 1][0].lstrip())
                if len(tail) < 5 or len(head) < 5:
                    continue

                split_pattern = tail + "\n\n" + head
                if split_pattern in document_text:
                    continue  # already separated

                # DI may join with a space or a single newline
                for joiner in (" ", "\n"):
                    merged_pattern = tail + joiner + head
                    if merged_pattern in document_text:
                        document_text = document_text.replace(merged_pattern, split_pattern, 1)
                        patches_applied += 1
                        break

    doc.close()

    if patches_applied:
        print(f"  [PDF-paragraph-patch] Restored paragraph breaks at {patches_applied} location(s).")
    return document_text


# =========================================================
# DOCX Table Extraction (supplement DI output)
# =========================================================

def extract_tables_from_docx(docx_path: str) -> list:
    """Extract all tables (including nested) from a DOCX file as formatted text rows."""
    doc = Document(docx_path)
    tables_text = []

    def _process_table(table):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            tables_text.append("\n".join(rows))
        for row in table.rows:
            for cell in row.cells:
                for nested in cell.tables:
                    _process_table(nested)

    for table in doc.tables:
        _process_table(table)
    return tables_text


def append_missing_tables(document_text: str, docx_path: str) -> str:
    """Append tables from DOCX that were not captured by Document Intelligence."""
    if not docx_path.lower().endswith(".docx"):
        return document_text

    docx_tables = extract_tables_from_docx(docx_path)
    if not docx_tables:
        return document_text

    missing_tables = []
    for table_text in docx_tables:
        all_cell_values = set()
        for line in table_text.split("\n"):
            for cell in line.split("|"):
                cleaned = cell.strip()
                if len(cleaned) > 10:
                    all_cell_values.add(cleaned)
        if all_cell_values:
            found = sum(1 for c in all_cell_values if c in document_text)
            if found < len(all_cell_values) / 2:
                missing_tables.append(table_text)

    if missing_tables:
        table_section = (
            "\n\n=============================\n"
            "ADDITIONAL TABLE DATA FROM DOCUMENT\n"
            "=============================\n"
        )
        table_section += "\n\n".join(missing_tables)
        print(f"  Appended {len(missing_tables)} table(s) not found in DI output.")
        return document_text + table_section

    return document_text


def append_structured_table_supplement(document_text: str, docx_path: str) -> str:
    """Append a structured table supplement with clear section/row/column markers.

    Document Intelligence flattens multi-column tables into a linear text stream,
    losing the column boundaries.  This function reads all tables from the DOCX
    using python-docx and appends them at the end with explicit row numbers and
    column headers so the LLM can correctly assign text to the right schema field.
    """
    if not docx_path.lower().endswith(".docx"):
        return document_text

    doc = Document(docx_path)
    structured_parts = []
    table_index = 0

    def _format_table(table, label_prefix=""):
        nonlocal table_index
        table_index += 1
        label = f"{label_prefix}Table {table_index}" if label_prefix else f"Table {table_index}"

        rows_data = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows_data.append(cells)

        if not rows_data:
            return

        # Use first row as column headers if it looks like a header
        col_headers = None
        data_start = 0
        if len(rows_data) > 1:
            first_row = rows_data[0]
            # Heuristic: header rows tend to have short text in most cells
            avg_len = sum(len(c) for c in first_row if c) / max(len(first_row), 1)
            if avg_len < 80:
                col_headers = first_row
                data_start = 1

        lines = [f"\n--- {label} ---"]
        if col_headers:
            lines.append(f"  Column headers: {' | '.join(col_headers)}")

        for row_idx, row_cells in enumerate(rows_data[data_start:], start=1):
            if col_headers:
                cell_parts = []
                for col_idx, cell_val in enumerate(row_cells):
                    header = col_headers[col_idx] if col_idx < len(col_headers) else f"Col{col_idx+1}"
                    cell_parts.append(f"    [{header}]: {cell_val if cell_val else '(empty)'}")
                lines.append(f"  Row {row_idx}:")
                lines.extend(cell_parts)
            else:
                lines.append(f"  Row {row_idx}: {' | '.join(row_cells)}")

        structured_parts.append("\n".join(lines))

        # Handle nested tables
        for row in table.rows:
            for cell in row.cells:
                for nested in cell.tables:
                    _format_table(nested, label_prefix=f"{label} > ")

    for table in doc.tables:
        _format_table(table)

    if structured_parts:
        supplement = (
            "\n\n=============================\n"
            "STRUCTURED TABLE DATA (with row numbers and column labels)\n"
            "Use this section to resolve ambiguities about which text belongs to which table column.\n"
            "=============================\n"
        )
        supplement += "\n".join(structured_parts)
        print(f"  [Table-supplement] Appended structured data for {table_index} table(s).")
        return document_text + supplement

    return document_text


def append_structured_table_supplement_pdf(document_text: str, pdf_path: str) -> str:
    """Append a structured table supplement for PDFs using PyMuPDF's table detection.

    This mirrors what append_structured_table_supplement does for DOCX files via
    python-docx.  PyMuPDF's page.find_tables() detects tables independently of
    Azure Document Intelligence, ensuring provision tables are captured even when
    DI flattens them into paragraph text.
    """
    if not pdf_path.lower().endswith(".pdf"):
        return document_text

    try:
        doc = fitz.open(pdf_path)
    except Exception as e:
        print(f"  [PDF-Table-supplement] Failed to open PDF: {e}")
        return document_text

    structured_parts = []
    table_index = 0

    for page_num in range(len(doc)):
        page = doc[page_num]
        try:
            tables = page.find_tables()
        except Exception:
            continue

        for table in tables:
            table_index += 1
            try:
                extracted = table.extract()
            except Exception:
                continue

            if not extracted or len(extracted) == 0:
                continue

            # Filter out empty rows
            rows_data = []
            for row in extracted:
                cells = [str(cell).strip() if cell else "" for cell in row]
                if any(cells):
                    rows_data.append(cells)

            if not rows_data:
                continue

            # Use first row as column headers if it looks like a header
            col_headers = None
            data_start = 0
            if len(rows_data) > 1:
                first_row = rows_data[0]
                avg_len = sum(len(c) for c in first_row if c) / max(len(first_row), 1)
                if avg_len < 80:
                    col_headers = first_row
                    data_start = 1

            lines = [f"\n--- Table {table_index} ---"]
            if col_headers:
                lines.append(f"  Column headers: {' | '.join(col_headers)}")

            for row_idx, row_cells in enumerate(rows_data[data_start:], start=1):
                if col_headers:
                    cell_parts = []
                    for col_idx, cell_val in enumerate(row_cells):
                        header = col_headers[col_idx] if col_idx < len(col_headers) else f"Col{col_idx+1}"
                        cell_parts.append(f"    [{header}]: {cell_val if cell_val else '(empty)'}")
                    lines.append(f"  Row {row_idx}:")
                    lines.extend(cell_parts)
                else:
                    lines.append(f"  Row {row_idx}: {' | '.join(row_cells)}")

            structured_parts.append("\n".join(lines))

    doc.close()

    if structured_parts:
        supplement = (
            "\n\n=============================\n"
            "STRUCTURED TABLE DATA (with row numbers and column labels)\n"
            "Use this section to resolve ambiguities about which text belongs to which table column.\n"
            "=============================\n"
        )
        supplement += "\n".join(structured_parts)
        print(f"  [PDF-Table-supplement] Appended structured data for {table_index} table(s).")
        return document_text + supplement

    return document_text


# =========================================================
# Full Document Text Builder (DI + hyperlinks + tables)
# =========================================================

def build_document_text(file_path: str) -> str:
    """Extract document text via DI, with hyperlink and table supplements."""
    print(f"  [DI] Analyzing: {file_path}")
    started_at = time.perf_counter()

    ext = os.path.splitext(file_path)[1].lower()

    try:
        doc_text = analyze_with_doc_intelligence(file_path)
        elapsed = time.perf_counter() - started_at
        print(f"  [DI] Done: {len(doc_text)} chars in {elapsed:.2f}s")
    except Exception as exc:
        if ext not in (".docx", ".pdf"):
            raise
        print(f"  [DI] Failed ({type(exc).__name__}), falling back to local extraction")
        doc_text = extract_text_locally(file_path)
        print(f"  [Local] Extracted {len(doc_text)} chars")

    # DOCX-specific supplements (hyperlinks, paragraph breaks, and table extraction)
    if ext == ".docx":
        doc_text = append_hyperlink_section(doc_text, file_path)
        doc_text = patch_table_paragraph_breaks(doc_text, file_path)
        doc_text = append_missing_tables(doc_text, file_path)
        doc_text = append_structured_table_supplement(doc_text, file_path)

    # PDF-specific supplements (paragraph breaks + structured table data via PyMuPDF)
    if ext == ".pdf":
        doc_text = patch_pdf_paragraph_breaks(doc_text, file_path)
        doc_text = append_structured_table_supplement_pdf(doc_text, file_path)

    doc_text = _repair_broken_urls(doc_text)
    return doc_text





# =========================================================
# LLM Extraction (deterministic, no agent)
# =========================================================

async def run_llm_extraction(full_prompt: str) -> tuple[str, dict]:
    """Send a single extraction prompt to the LLM and return the response."""
    client = create_openai_client()

    response = await asyncio.wait_for(
        client.chat.completions.create(
            model=deployment,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are an EHCP document extraction agent. "
                        "Return ONLY valid JSON matching the provided schema. "
                        "Do not include explanations, markdown, or code fences. "
                        "Preserve original wording from the document. "
                        "Do not invent or fabricate information. "
                        "CRITICAL: Extract ALL content from EVERY section completely. "
                        "Do NOT truncate, summarise, or skip any lines, paragraphs, or bullet points. "
                        "Every line of text that belongs to a schema field MUST appear in the output. "
                        "If a section has an introductory sentence followed by multiple lines or paragraphs, "
                        "include ALL of them — not just the first line."
                    ),
                },
                {"role": "user", "content": full_prompt},
            ],
            max_completion_tokens=32768,
            temperature=0.0,
            seed=42,
            top_p=1.0,
        ),
        timeout=LLM_TIMEOUT_SECONDS,
    )
    # Extract token usage if available
    token_usage = {}
    if hasattr(response, "usage") and response.usage:
        token_usage = {
            "prompt_tokens": response.usage.prompt_tokens or 0,
            "completion_tokens": response.usage.completion_tokens or 0,
            "total_tokens": response.usage.total_tokens or 0,
        }
    _accumulate_tokens(token_usage)
    return response.choices[0].message.content, token_usage


def split_schema(schema_str: str) -> list:
    """Split schema into chunks for extraction. Each chunk gets its own LLM call.
    
    For schemas with many keys (>3), each key gets its own chunk to ensure
    the LLM focuses on extracting ALL data for that section without losing content.
    Small trailing chunks are merged into the previous one.
    Related provision sections (H1/H2) are kept together so the LLM can
    correctly partition rows between them.
    """
    schema = json.loads(schema_str)
    keys = list(schema.keys())
    n = len(keys)
    if n <= 3:
        return [schema]

    # Group related keys that must be extracted together
    _PROVISION_GROUPS = {
        "H1_Social_Care_Provision", "H2_Social_Care_Provision",
    }

    chunks = []
    grouped_keys = set()

    # First pass: create grouped chunks for related keys
    for key in keys:
        if key in grouped_keys:
            continue
        if key in _PROVISION_GROUPS:
            # Combine all provision keys into one chunk
            group_chunk = {}
            for gk in keys:
                if gk in _PROVISION_GROUPS:
                    group_chunk[gk] = schema[gk]
                    grouped_keys.add(gk)
            chunks.append(group_chunk)
        else:
            chunks.append({key: schema[key]})
            grouped_keys.add(key)

    # Merge last chunk into previous if it has a simple scalar value (e.g. Advice_Giver)
    if len(chunks) > 1:
        last_val = list(chunks[-1].values())[0]
        if not isinstance(last_val, dict) or len(last_val) <= 3:
            last = chunks.pop()
            chunks[-1].update(last)
    return chunks


async def extract_document(document_text: str, prompt_file: str, schema_file: str, chunked: bool = False) -> dict:
    """Run extraction on a single document (single-pass or chunked)."""
    prompt_template = load_text_file(prompt_file)
    schema_str = load_schema_json(schema_file)
    schema = json.loads(schema_str)
    schema_keys = list(schema.keys())

    use_chunked = chunked or (len(document_text) > 10000 and len(schema_keys) > 3)

    if use_chunked:
        schema_chunks = split_schema(schema_str)
        print(f"  Using chunked extraction ({len(document_text)} chars, {len(schema_chunks)} passes)")
        data = {}
        for i, chunk in enumerate(schema_chunks):
            section_names = ", ".join(chunk.keys())
            chunk_prompt = prompt_template.replace(
                "{schema}", json.dumps(chunk, indent=2)
            ).replace(
                "{document_text}", document_text
            )
            chunk_prompt += (
                f"\n\nIMPORTANT: Extract ONLY these sections: {section_names}. "
                "Ignore all other sections.\n"
                "COMPLETENESS RULES:\n"
                "- Include ALL text from EVERY field — do NOT truncate or shorten any content.\n"
                "- If a section has an introductory line followed by multiple lines, paragraphs or items, "
                "extract ALL of them as separate array elements.\n"
                "- Do NOT return only the first line of a multi-line section.\n"
                "- Every line in the document that falls under a schema field MUST appear in your output.\n"
                "- For provision table cells, include the COMPLETE cell text with all sub-items.\n"
                "- Count the items you extract and verify nothing was skipped.\n"
                "- For H1/H2 provision tables: use the STRUCTURED TABLE DATA at the end to count exact rows. "
                "All four arrays (Provision, By_Whom, Frequency, Met_Need) MUST have the SAME length. "
                "If a Met_Need cell has multiple lines, join them with \\n into ONE string — do NOT "
                "split them into separate Provision entries."
            )
            print(f"  Pass {i+1}/{len(schema_chunks)} ({section_names})...")
            result_text, _usage = await run_llm_extraction(chunk_prompt)
            chunk_data = safe_parse_json(result_text)
            print(f"  Pass {i+1} done. Keys: {list(chunk_data.keys())}")
            data.update(chunk_data)
        return post_process_json(data)
    else:
        full_prompt = prompt_template.replace(
            "{schema}", schema_str
        ).replace(
            "{document_text}", document_text
        )
        full_prompt += (
            "\n\nCOMPLETENESS RULES:\n"
            "- Include ALL text from EVERY field — do NOT truncate or shorten any content.\n"
            "- If a section has an introductory line followed by multiple lines, paragraphs or items, "
            "extract ALL of them as separate array elements.\n"
            "- Do NOT return only the first line of a multi-line section.\n"
            "- Every line in the document that falls under a schema field MUST appear in your output.\n"
            "- Do NOT return null for fields that have data in the document."
        )
        print(f"  Using single-pass extraction ({len(document_text)} chars)")
        result_text, _usage = await run_llm_extraction(full_prompt)
        return post_process_json(safe_parse_json(result_text))
