#python validation_code.py --docx EHCP_LCC_Template.docx --input-folder input/Henry --mapping-workbook ehcp_mapping.xlsx
import argparse
import json
import os
import re
import sys
from dataclasses import dataclass, field
from typing import Any, Dict, Iterable, Optional

import openpyxl
from docx import Document


PERSONAL_DETAILS_SHEET = "PersonalDetails"
EDUCATION_DETAILS_SHEET = "EducationAdvice"
HEALTH_DETAILS_SHEET = "HealthAdvice"
SOCIAL_CARE_DETAILS_SHEET = "SocialCareAdvice"

PERSONAL_REQUIRED_KEYS = {
    "name",
    "preferred_name",
    "date_of_birth",
    "sex",
    "ethnicity",
    "main_contact_parent_or_carer_1",
    "main_contact_relationship",
    "main_contact_email",
    "main_contact_telephone_number",
}

EDUCATION_EXPECTED_SECTIONS = {
    "Pupil_Views",
    "Parent_Carer_Views",
    "Cognition_and_Learning",
    "Communication_and_Interaction",
    "Social_Emotional_Mental_Health",
    "Sensory_and_Physical",
}

SOCIALCARE_EXPECTED_SECTIONS = {
    "Social_Care_Needs",
    "Outcomes",
    "H1_Social_Care_Provision",
    "H2_Social_Care_Provision",
    "Advice_Giver",
}

HEALTH_EXPECTED_KEYS = {
    "Child_Details",
    "Medical_History",
    "Health_Needs",
    "Outcomes",
    "Health_Provision",
    "Advice_Date",
}

WORKBOOK_HEADERS = {
    PERSONAL_DETAILS_SHEET: ("row_group", "column_order", "label", "json_key"),
    EDUCATION_DETAILS_SHEET: ("Template column", "Template section", "key", "section", "Extra details"),
    HEALTH_DETAILS_SHEET: ("Template column", "Template section", "key", "section", "Extra details"),
    SOCIAL_CARE_DETAILS_SHEET: ("Template column", "Template section", "key", "section", "Extra details"),
}

TABLE_HEADERS = ("Check", "Target", "Status", "Details")
ISSUE_STATUSES = {"Missing", "Invalid", "Mismatch", "Type mismatch", "Incomplete"}


@dataclass
class ValidationTableRow:
    check: str
    target: str
    status: str
    details: str = ""


@dataclass
class ValidationReport:
    errors: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    infos: list[str] = field(default_factory=list)
    tables: Dict[str, list[ValidationTableRow]] = field(default_factory=dict)

    def error(self, message: str) -> None:
        self.errors.append(message)

    def warn(self, message: str) -> None:
        self.warnings.append(message)

    def info(self, message: str) -> None:
        self.infos.append(message)

    def add_table_row(self, table_name: str, check: str, target: str, status: str, details: str = "") -> None:
        self.tables.setdefault(table_name, []).append(
            ValidationTableRow(check=check, target=target, status=status, details=details)
        )

    @property
    def is_valid(self) -> bool:
        return not self.errors


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Validate EHCP input JSON files and mapping workbook before running writer_Code.py."
    )
    parser.add_argument("--docx", required=True, help="Path to source DOCX file")
    parser.add_argument(
        "--input-folder",
        default=None,
        help="Path to a child input folder containing personal_details.json, education_advice.json, socialcare_advice.json, and health_advice.json",
    )
    parser.add_argument("--json", default=os.path.join("input", "personal_details.json"), help="Path to personal details JSON file")
    parser.add_argument(
        "--education-json",
        default=os.path.join("input", "education_advice.json"),
        help="Path to extracted education JSON file",
    )
    parser.add_argument(
        "--socialcare-json",
        default=os.path.join("input", "socialcare_advice.json"),
        help="Path to social care advice JSON file",
    )
    parser.add_argument(
        "--health-json",
        default=os.path.join("input", "health_advice.json"),
        help="Path to health advice JSON file",
    )
    parser.add_argument(
        "--mapping-workbook",
        default="ehcp_mapping.xlsx",
        help="Path to workbook with mapping sheets for personal, education, health, and social care",
    )
    parser.add_argument("--education-mapping", default=None, help="Deprecated alias for mapping workbook path")
    parser.add_argument(
        "--report-file",
        default=None,
        help="Optional path to save a Markdown validation report with tabular results",
    )
    parser.add_argument(
        "--json-report-file",
        default=None,
        help="Optional path to save the validation report as JSON",
    )
    parser.add_argument(
        "--word-report-file",
        default=None,
        help="Optional path to save the validation report as a Word document (.docx)",
    )
    parser.add_argument(
        "--actual-output-docx",
        default=None,
        help="Optional path to the generated filled DOCX to compare against an expected target DOCX",
    )
    parser.add_argument(
        "--expected-output-docx",
        default=None,
        help="Optional path to the expected target DOCX for output comparison",
    )
    return parser.parse_args()


def find_case_insensitive_match(folder_path: str, file_name: str) -> str:
    exact_path = os.path.join(folder_path, file_name)
    if os.path.exists(exact_path):
        return exact_path

    if not os.path.isdir(folder_path):
        return exact_path

    target = file_name.lower()
    for entry in os.listdir(folder_path):
        if entry.lower() == target:
            return os.path.join(folder_path, entry)
    return exact_path


def resolve_input_json_paths(args: argparse.Namespace) -> Dict[str, str]:
    if not args.input_folder:
        return {
            "json": args.json,
            "education_json": args.education_json,
            "socialcare_json": args.socialcare_json,
            "health_json": args.health_json,
        }

    return {
        "json": find_case_insensitive_match(args.input_folder, "personal_details.json"),
        "education_json": find_case_insensitive_match(args.input_folder, "education_advice.json"),
        "socialcare_json": find_case_insensitive_match(args.input_folder, "socialcare_advice.json"),
        "health_json": find_case_insensitive_match(args.input_folder, "health_advice.json"),
    }


def choose_mapping_workbook(args: argparse.Namespace) -> str:
    if os.path.exists(args.mapping_workbook):
        return args.mapping_workbook
    if args.education_mapping and os.path.exists(args.education_mapping):
        return args.education_mapping
    return args.mapping_workbook


def load_json_file(file_path: str, report: ValidationReport, label: str, required: bool) -> Optional[Dict[str, Any]]:
    if not os.path.exists(file_path):
        message = f"{label} file not found: {file_path}"
        if required:
            report.error(message)
        else:
            report.warn(message)
        return None

    try:
        with open(file_path, "r", encoding="utf-8") as file_handle:
            content = json.load(file_handle)
    except json.JSONDecodeError as exc:
        report.error(f"{label} contains invalid JSON at line {exc.lineno}, column {exc.colno}: {exc.msg}")
        return None
    except OSError as exc:
        report.error(f"Unable to read {label}: {exc}")
        return None

    if not isinstance(content, dict):
        report.error(f"{label} must contain a top-level JSON object")
        return None

    report.info(f"Loaded {label}: {file_path}")
    return content


def validate_docx(docx_path: str, report: ValidationReport) -> None:
    if not os.path.exists(docx_path):
        report.error(f"DOCX file not found: {docx_path}")
        return
    if not docx_path.lower().endswith((".docx", ".pdf")):
        report.warn(f"File path does not end with .docx or .pdf: {docx_path}")
    report.info(f"Found DOCX template: {docx_path}")


def normalize_doc_value(value: str) -> str:
    text = str(value or "")
    text = text.replace("\r\n", "\n")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def read_docx_snapshot(docx_path: str) -> Dict[str, Any]:
    document = Document(docx_path)
    paragraphs = [normalize_doc_value(paragraph.text) for paragraph in document.paragraphs]
    paragraphs = [paragraph for paragraph in paragraphs if paragraph]

    tables = []
    for table_index, table in enumerate(document.tables):
        rows = []
        for row_index, row in enumerate(table.rows):
            cells = []
            for cell_index, cell in enumerate(row.cells):
                cell_text = normalize_doc_value(cell.text)
                cells.append(
                    {
                        "table_index": table_index,
                        "row_index": row_index,
                        "cell_index": cell_index,
                        "text": cell_text,
                    }
                )
            rows.append(cells)
        tables.append(rows)

    return {"paragraphs": paragraphs, "tables": tables}


def compare_docx_outputs(actual_docx_path: str, expected_docx_path: str, report: ValidationReport) -> None:
    if not os.path.exists(actual_docx_path):
        report.error(f"Generated output DOCX not found: {actual_docx_path}")
        report.add_table_row("Output Document Comparison", "Generated file", actual_docx_path, "Missing", "Generated output file not found")
        return

    if not os.path.exists(expected_docx_path):
        report.error(f"Expected target DOCX not found: {expected_docx_path}")
        report.add_table_row("Output Document Comparison", "Expected file", expected_docx_path, "Missing", "Expected target file not found")
        return

    report.info(f"Comparing generated DOCX against expected DOCX: {actual_docx_path} vs {expected_docx_path}")

    actual_snapshot = read_docx_snapshot(actual_docx_path)
    expected_snapshot = read_docx_snapshot(expected_docx_path)

    actual_paragraphs = actual_snapshot["paragraphs"]
    expected_paragraphs = expected_snapshot["paragraphs"]
    report.add_table_row(
        "Output Document Comparison",
        "Paragraph count",
        "Document paragraphs",
        "Match" if len(actual_paragraphs) == len(expected_paragraphs) else "Mismatch",
        f"generated={len(actual_paragraphs)}; expected={len(expected_paragraphs)}",
    )

    max_paragraphs = max(len(actual_paragraphs), len(expected_paragraphs))
    paragraph_differences = 0
    for index in range(max_paragraphs):
        actual_text = actual_paragraphs[index] if index < len(actual_paragraphs) else ""
        expected_text = expected_paragraphs[index] if index < len(expected_paragraphs) else ""
        if actual_text != expected_text:
            paragraph_differences += 1
            report.add_table_row(
                "Output Document Comparison",
                "Paragraph content",
                f"Paragraph {index + 1}",
                "Mismatch",
                f"expected='{expected_text[:120]}' ; generated='{actual_text[:120]}'",
            )
            if paragraph_differences >= 10:
                break

    if paragraph_differences == 0:
        report.add_table_row(
            "Output Document Comparison",
            "Paragraph content",
            "Paragraph text",
            "Match",
            "Paragraph content matches the expected document",
        )

    actual_tables = actual_snapshot["tables"]
    expected_tables = expected_snapshot["tables"]
    report.add_table_row(
        "Output Document Comparison",
        "Table count",
        "Document tables",
        "Match" if len(actual_tables) == len(expected_tables) else "Mismatch",
        f"generated={len(actual_tables)}; expected={len(expected_tables)}",
    )

    max_tables = max(len(actual_tables), len(expected_tables))
    cell_differences = 0
    for table_index in range(max_tables):
        actual_table = actual_tables[table_index] if table_index < len(actual_tables) else []
        expected_table = expected_tables[table_index] if table_index < len(expected_tables) else []
        max_rows = max(len(actual_table), len(expected_table))

        for row_index in range(max_rows):
            actual_row = actual_table[row_index] if row_index < len(actual_table) else []
            expected_row = expected_table[row_index] if row_index < len(expected_table) else []
            max_cells = max(len(actual_row), len(expected_row))

            for cell_index in range(max_cells):
                actual_text = actual_row[cell_index]["text"] if cell_index < len(actual_row) else ""
                expected_text = expected_row[cell_index]["text"] if cell_index < len(expected_row) else ""
                if actual_text != expected_text:
                    cell_differences += 1
                    report.add_table_row(
                        "Output Document Comparison",
                        "Table cell",
                        f"Table {table_index + 1}, Row {row_index + 1}, Cell {cell_index + 1}",
                        "Mismatch",
                        f"expected='{expected_text[:120]}' ; generated='{actual_text[:120]}'",
                    )
                    if cell_differences >= 20:
                        break
            if cell_differences >= 20:
                break
        if cell_differences >= 20:
            break

    if cell_differences == 0:
        report.add_table_row(
            "Output Document Comparison",
            "Table cell",
            "Table content",
            "Match",
            "Table content matches the expected document",
        )


def validate_personal_json(data: Optional[Dict[str, Any]], report: ValidationReport) -> None:
    if data is None:
        return

    for key in sorted(PERSONAL_REQUIRED_KEYS):
        if key in data:
            value = data[key]
            if value is None or (isinstance(value, str) and not value.strip()):
                report.add_table_row("Personal Details", "Expected key", key, "Missing",
                                     "Key exists but value is empty/null — source document may be missing this information")
            else:
                report.add_table_row("Personal Details", "Expected key", key, "Present", f"type={type(value).__name__}")
        else:
            report.add_table_row("Personal Details", "Expected key", key, "Missing", "Key not found in personal details JSON")

    missing_keys = sorted(key for key in PERSONAL_REQUIRED_KEYS if key not in data)
    if missing_keys:
        report.warn("Personal details JSON is missing expected keys: " + ", ".join(missing_keys))

    # Check for empty required values (not just missing keys)
    empty_required = sorted(
        key for key in PERSONAL_REQUIRED_KEYS
        if key in data and (data[key] is None or (isinstance(data[key], str) and not data[key].strip()))
    )
    if empty_required:
        report.warn("Personal details JSON has empty values for required fields: " + ", ".join(empty_required))

    if not (str(data.get("name") or "").strip() or str(data.get("preferred_name") or "").strip()):
        report.error("Personal details JSON must include at least one of 'name' or 'preferred_name'")
        report.add_table_row("Personal Details", "Critical field", "name/preferred_name", "Invalid", "Both values are empty")
    else:
        report.add_table_row("Personal Details", "Critical field", "name/preferred_name", "Present", "At least one value is populated")


def _is_section_empty(value) -> bool:
    """Check if a section value has no substantive content."""
    if value is None:
        return True
    if isinstance(value, str):
        return not value.strip()
    if isinstance(value, list):
        return not any(
            (isinstance(item, str) and item.strip()) or
            (isinstance(item, dict) and not _is_section_empty(item)) or
            (isinstance(item, list) and not _is_section_empty(item))
            for item in value
        )
    if isinstance(value, dict):
        return all(_is_section_empty(v) for v in value.values())
    return False


def validate_section_types(
    data: Dict[str, Any],
    expected_sections: Iterable[str],
    report: ValidationReport,
    label: str,
    table_name: str,
) -> None:
    missing_sections = sorted(section for section in expected_sections if section not in data)
    if missing_sections:
        report.warn(f"{label} is missing expected sections: " + ", ".join(missing_sections))

    empty_sections = []
    for section in expected_sections:
        if section not in data:
            report.add_table_row(table_name, "Expected section", section, "Missing", "Section not found")
            continue

        section_value = data[section]
        section_type = type(section_value).__name__
        if section_value is not None and not isinstance(section_value, (dict, list, str)):
            report.warn(f"{label} section '{section}' has unexpected type: {section_type}")
            report.add_table_row(table_name, "Expected section", section, "Type mismatch", f"type={section_type}")
        elif _is_section_empty(section_value):
            empty_sections.append(section)
            report.add_table_row(table_name, "Section content", section, "Missing",
                                 "Section exists but all fields are empty/null — source document may be missing this section's content")
        else:
            report.add_table_row(table_name, "Expected section", section, "Present", f"type={section_type}")

    if empty_sections:
        report.warn(f"{label} has sections with no substantive content: " + ", ".join(empty_sections))


def validate_education_json(data: Optional[Dict[str, Any]], report: ValidationReport) -> None:
    if data is None:
        return
    validate_section_types(data, EDUCATION_EXPECTED_SECTIONS, report, "Education advice JSON", "Education Advice")

    # Check each education section for content completeness
    provision_sections = ["Cognition_and_Learning", "Communication_and_Interaction",
                          "Social_Emotional_Mental_Health", "Sensory_and_Physical"]
    provision_fields = ["Strengths", "Needs", "Provision"]

    for section_name in provision_sections:
        section = data.get(section_name)
        if not isinstance(section, dict):
            continue
        for field_name in provision_fields:
            value = section.get(field_name)
            if _is_section_empty(value):
                report.add_table_row("Education Advice", "Content completeness",
                                     f"{section_name}.{field_name}", "Missing",
                                     f"Expected content in {field_name} but field is empty/null")

    for section in EDUCATION_EXPECTED_SECTIONS:
        value = data.get(section)
        if value is not None and not isinstance(value, dict):
            report.warn(f"Education section '{section}' should usually be an object, found {type(value).__name__}")


def validate_socialcare_json(data: Optional[Dict[str, Any]], report: ValidationReport) -> None:
    if data is None:
        return
    validate_section_types(data, SOCIALCARE_EXPECTED_SECTIONS, report, "Social care advice JSON", "Social Care Advice")

    advice_giver = data.get("Advice_Giver")
    if advice_giver is not None and not isinstance(advice_giver, dict):
        report.warn("Social care Advice_Giver should be an object with Name, Role, and Date")

    # Check provision table content
    for prov_section in ["H1_Social_Care_Provision", "H2_Social_Care_Provision"]:
        section = data.get(prov_section)
        if isinstance(section, dict):
            for field_name in ["Provision", "By_Whom"]:
                value = section.get(field_name)
                if _is_section_empty(value):
                    report.add_table_row("Social Care Advice", "Content completeness",
                                         f"{prov_section}.{field_name}", "Missing",
                                         f"Expected content in {field_name} but field is empty/null")


def validate_health_json(data: Optional[Dict[str, Any]], report: ValidationReport) -> None:
    if data is None:
        return

    for key in sorted(HEALTH_EXPECTED_KEYS):
        if key in data:
            value = data[key]
            if _is_section_empty(value):
                report.add_table_row("Health Advice", "Expected key", key, "Missing",
                                     "Key exists but value is empty/null — source document may be missing this content")
            else:
                report.add_table_row("Health Advice", "Expected key", key, "Present", f"type={type(value).__name__}")
        else:
            report.add_table_row("Health Advice", "Expected key", key, "Missing", "Key not found in health advice JSON")

    missing_keys = sorted(key for key in HEALTH_EXPECTED_KEYS if key not in data)
    if missing_keys:
        report.warn("Health advice JSON is missing expected keys: " + ", ".join(missing_keys))

    # Check for empty values in expected keys
    empty_keys = sorted(
        key for key in HEALTH_EXPECTED_KEYS
        if key in data and _is_section_empty(data[key])
    )
    if empty_keys:
        report.warn("Health advice JSON has empty values for expected fields: " + ", ".join(empty_keys))

    if "Name" not in data or "Role" not in data:
        report.warn(
            "Health advice JSON does not include top-level 'Name' and 'Role'. writer_Code.py currently reads those keys when filling section K."
        )
        report.add_table_row("Health Advice", "Writer compatibility", "Name/Role", "Missing", "writer_Code.py expects top-level Name and Role for section K")
    else:
        report.add_table_row("Health Advice", "Writer compatibility", "Name/Role", "Present", "Top-level Name and Role exist")


def load_workbook_safe(workbook_path: str, report: ValidationReport) -> Optional[openpyxl.Workbook]:
    if not os.path.exists(workbook_path):
        report.error(f"Mapping workbook not found: {workbook_path}")
        return None
    try:
        workbook = openpyxl.load_workbook(workbook_path, data_only=True)
    except Exception as exc:
        report.error(f"Unable to open mapping workbook '{workbook_path}': {exc}")
        return None

    report.info(f"Loaded mapping workbook: {workbook_path}")
    return workbook


def validate_workbook_headers(workbook: openpyxl.Workbook, report: ValidationReport) -> None:
    for sheet_name, expected_header in WORKBOOK_HEADERS.items():
        if sheet_name not in workbook.sheetnames:
            report.error(f"Workbook is missing required sheet: {sheet_name}")
            report.add_table_row("Workbook", "Required sheet", sheet_name, "Missing", "Sheet not found in workbook")
            continue

        sheet = workbook[sheet_name]
        actual_header = next(sheet.iter_rows(values_only=True), ())
        actual_tuple = tuple(actual_header[: len(expected_header)])
        if actual_tuple != expected_header:
            report.error(
                f"Sheet '{sheet_name}' has unexpected header. Expected {expected_header}, found {actual_tuple}"
            )
            report.add_table_row("Workbook", "Sheet header", sheet_name, "Mismatch", f"expected={expected_header}; found={actual_tuple}")
        else:
            report.add_table_row("Workbook", "Sheet header", sheet_name, "Match", "Required sheet and header found")


def iter_mapping_rows(sheet) -> Iterable[tuple[int, tuple[Any, ...]]]:
    for row_index, row in enumerate(sheet.iter_rows(values_only=True), start=1):
        if row_index == 1:
            continue
        yield row_index, row


def validate_personal_mapping(
    workbook: openpyxl.Workbook,
    personal_data: Optional[Dict[str, Any]],
    report: ValidationReport,
) -> None:
    if personal_data is None or PERSONAL_DETAILS_SHEET not in workbook.sheetnames:
        return

    sheet = workbook[PERSONAL_DETAILS_SHEET]
    for row_index, row in iter_mapping_rows(sheet):
        label = row[2] if len(row) > 2 else None
        json_key = row[3] if len(row) > 3 else None
        if not label and not json_key:
            continue
        if not label or not json_key:
            report.warn(f"{PERSONAL_DETAILS_SHEET} row {row_index} is missing label or json_key")
            report.add_table_row("Personal Details", "Mapping row", f"row {row_index}", "Incomplete", "label or json_key is missing")
            continue
        if str(json_key).strip() not in personal_data:
            report.warn(
                f"{PERSONAL_DETAILS_SHEET} row {row_index} references missing personal JSON key '{str(json_key).strip()}'"
            )
            report.add_table_row(
                "Personal Details",
                "Mapping key",
                str(json_key).strip(),
                "Missing",
                f"Referenced by {PERSONAL_DETAILS_SHEET} row {row_index}",
            )
        else:
            report.add_table_row(
                "Personal Details",
                "Mapping key",
                str(json_key).strip(),
                "Match",
                f"Referenced by {PERSONAL_DETAILS_SHEET} row {row_index}",
            )


def lookup_mapping_target(data: Dict[str, Any], section_name: str, key_field: str) -> tuple[bool, bool]:
    if section_name:
        if section_name not in data:
            return False, False
        source = data.get(section_name)
    else:
        source = data

    if not isinstance(source, dict):
        return True, False

    key_names = [item.strip() for item in str(key_field or "").split(",") if item.strip()]
    if not key_names:
        return True, True

    return True, all(key_name in source for key_name in key_names)


def validate_section_mapping(
    workbook: openpyxl.Workbook,
    sheet_name: str,
    data: Optional[Dict[str, Any]],
    report: ValidationReport,
) -> None:
    if data is None or sheet_name not in workbook.sheetnames:
        return

    sheet = workbook[sheet_name]
    table_name = {
        EDUCATION_DETAILS_SHEET: "Education Advice",
        SOCIAL_CARE_DETAILS_SHEET: "Social Care Advice",
        HEALTH_DETAILS_SHEET: "Health Advice",
    }.get(sheet_name, sheet_name)

    for row_index, row in iter_mapping_rows(sheet):
        template_column = str(row[0] or "").strip() if len(row) > 0 else ""
        template_section = str(row[1] or "").strip() if len(row) > 1 else ""
        key_field = str(row[2] or "").strip() if len(row) > 2 else ""
        section_name = str(row[3] or "").strip() if len(row) > 3 else ""

        if not any([template_column, template_section, key_field, section_name]):
            continue

        if not key_field:
            report.warn(f"{sheet_name} row {row_index} is missing the key column")
            report.add_table_row(table_name, "Mapping row", f"row {row_index}", "Incomplete", "key column is missing")
            continue

        section_exists, keys_exist = lookup_mapping_target(data, section_name, key_field)
        if not section_exists:
            location = section_name or "<root>"
            report.warn(
                f"{sheet_name} row {row_index} references missing section '{location}' for template section '{template_section}'"
            )
            report.add_table_row(
                table_name,
                "Mapping section",
                f"{location} -> {key_field}",
                "Missing",
                f"Referenced by row {row_index} for template section '{template_section}'",
            )
            continue

        if not keys_exist:
            report.warn(
                f"{sheet_name} row {row_index} references missing key(s) '{key_field}' in section '{section_name or '<root>'}'"
            )
            report.add_table_row(
                table_name,
                "Mapping key",
                f"{section_name or '<root>'} -> {key_field}",
                "Missing",
                f"Referenced by row {row_index}",
            )
        else:
            report.add_table_row(
                table_name,
                "Mapping key",
                f"{section_name or '<root>'} -> {key_field}",
                "Match",
                f"Referenced by row {row_index}",
            )


def format_ascii_table(rows: list[ValidationTableRow]) -> list[str]:
    widths = [len(header) for header in TABLE_HEADERS]
    cell_rows = [(row.check, row.target, row.status, row.details) for row in rows]

    for cell_row in cell_rows:
        for index, value in enumerate(cell_row):
            widths[index] = min(max(widths[index], len(str(value))), 80)

    def build_border() -> str:
        return "+" + "+".join("-" * (width + 2) for width in widths) + "+"

    def build_row(values: tuple[str, str, str, str]) -> str:
        padded = []
        for index, value in enumerate(values):
            text = str(value)
            if len(text) > widths[index]:
                text = text[: widths[index] - 3] + "..."
            padded.append(" " + text.ljust(widths[index]) + " ")
        return "|" + "|".join(padded) + "|"

    lines = [build_border(), build_row(TABLE_HEADERS), build_border()]
    for cell_row in cell_rows:
        lines.append(build_row(cell_row))
    lines.append(build_border())
    return lines


def prettify_text(value: str) -> str:
    text = str(value or "").replace("<root>", "Top level")
    text = text.replace("->", " to ")
    text = text.replace("_", " ")
    text = " ".join(text.split())
    return text[:1].upper() + text[1:] if text else ""


def display_status(status: str) -> str:
    mapping = {
        "Present": "OK",
        "Match": "OK",
        "Missing": "Missing",
        "Invalid": "Needs review",
        "Mismatch": "Needs review",
        "Type mismatch": "Needs review",
        "Incomplete": "Needs review",
    }
    return mapping.get(status, status)


def summarize_rows(rows: list[ValidationTableRow]) -> Dict[str, Any]:
    total = len(rows)
    issue_rows = [row for row in rows if row.status in ISSUE_STATUSES]
    ok_rows = [row for row in rows if row.status not in ISSUE_STATUSES]
    return {
        "total_checks": total,
        "ok_count": len(ok_rows),
        "issue_count": len(issue_rows),
        "result": "Needs attention" if issue_rows else "Ready",
        "issue_rows": issue_rows,
        "ok_rows": ok_rows,
    }


def build_issue_message(row: ValidationTableRow) -> str:
    target = prettify_text(row.target)
    details = row.details.strip()

    if row.check == "Expected key" and row.status == "Missing":
        return f"Expected field is missing: {target}."
    if row.check == "Expected section" and row.status == "Missing":
        return f"Expected section is missing: {target}."
    if row.check == "Mapping key" and row.status == "Missing":
        return f"A mapped field could not be found in the source data: {target}."
    if row.check == "Mapping section" and row.status == "Missing":
        return f"A mapped section could not be found in the source data: {target}."
    if row.check == "Writer compatibility" and row.status == "Missing":
        return f"The current document writer expects this information but it is missing: {target}."
    if row.check == "Sheet header" and row.status == "Mismatch":
        return f"Workbook sheet structure does not match what the validator expects: {target}."
    if row.check == "Required sheet" and row.status == "Missing":
        return f"Required workbook sheet is missing: {target}."
    if row.check == "Generated file" and row.status == "Missing":
        return f"The generated output document could not be found: {target}."
    if row.check == "Expected file" and row.status == "Missing":
        return f"The expected target document could not be found: {target}."
    if row.check == "Paragraph count" and row.status == "Mismatch":
        return f"The generated document has a different number of text paragraphs than the expected document: {row.details}."
    if row.check == "Paragraph content" and row.status == "Mismatch":
        return f"Some paragraph text does not match the expected output: {target}."
    if row.check == "Table count" and row.status == "Mismatch":
        return f"The generated document has a different number of tables than the expected document: {row.details}."
    if row.check == "Table cell" and row.status == "Mismatch":
        return f"Some table content does not match the expected output: {target}."
    if row.check == "Critical field" and row.status == "Invalid":
        return f"A critical value is empty or invalid: {target}."
    if row.check == "Mapping row" and row.status == "Incomplete":
        return f"A workbook mapping row is incomplete: {target}."

    if details:
        return f"{prettify_text(row.check)}: {target}. {details}"
    return f"{prettify_text(row.check)}: {target}."


def build_area_summary_rows(report: ValidationReport) -> list[tuple[str, str, str, str, str]]:
    summary_rows = []
    for table_name, rows in report.tables.items():
        summary = summarize_rows(rows)
        summary_rows.append(
            (
                table_name,
                summary["result"],
                str(summary["total_checks"]),
                str(summary["ok_count"]),
                str(summary["issue_count"]),
            )
        )
    return summary_rows


def build_summary_table_lines(report: ValidationReport) -> list[str]:
    headers = ("Area", "Result", "Checks", "OK", "Needs attention")
    rows = build_area_summary_rows(report)
    widths = [len(header) for header in headers]

    for row in rows:
        for index, value in enumerate(row):
            widths[index] = max(widths[index], len(value))

    def border() -> str:
        return "+" + "+".join("-" * (width + 2) for width in widths) + "+"

    def line(values: tuple[str, ...]) -> str:
        return "|" + "|".join(" " + values[index].ljust(widths[index]) + " " for index in range(len(values))) + "|"

    output = [border(), line(headers), border()]
    for row in rows:
        output.append(line(row))
    output.append(border())
    return output


def build_markdown_report(report: ValidationReport) -> str:
    lines = ["# Validation Report", ""]

    lines.append("## At A Glance")
    lines.append("")
    lines.append("| Area | Result | Checks | OK | Needs attention |")
    lines.append("| --- | --- | --- | --- | --- |")
    for area, result, checks, ok_count, issue_count in build_area_summary_rows(report):
        lines.append(f"| {area} | {result} | {checks} | {ok_count} | {issue_count} |")
    lines.append("")

    for table_name, rows in report.tables.items():
        summary = summarize_rows(rows)
        lines.append(f"## {table_name}")
        lines.append("")
        lines.append(f"Result: {summary['result']}")
        lines.append("")
        lines.append(f"- Total checks: {summary['total_checks']}")
        lines.append(f"- OK: {summary['ok_count']}")
        lines.append(f"- Needs attention: {summary['issue_count']}")
        lines.append("")

        lines.append("### What Needs Attention")
        lines.append("")
        if summary["issue_rows"]:
            for row in summary["issue_rows"]:
                lines.append(f"- {build_issue_message(row)}")
        else:
            lines.append("- No problems found in this section.")
        lines.append("")

        lines.append("### Detailed Check Table")
        lines.append("")
        lines.append("| What was checked | Item | Result | Explanation |")
        lines.append("| --- | --- | --- | --- |")
        for row in rows:
            details = row.details.replace("|", "\\|")
            target = prettify_text(row.target).replace("|", "\\|")
            check = prettify_text(row.check).replace("|", "\\|")
            status = display_status(row.status).replace("|", "\\|")
            lines.append(f"| {check} | {target} | {status} | {details} |")
        lines.append("")

    lines.append("## Summary")
    lines.append("")
    lines.append(f"- Errors: {len(report.errors)}")
    lines.append(f"- Warnings: {len(report.warnings)}")
    lines.append(f"- Status: {'FAILED' if report.errors else 'PASSED'}")
    lines.append("")

    if report.warnings:
        lines.append("## Warnings")
        lines.append("")
        for warning in report.warnings:
            lines.append(f"- {warning}")
        lines.append("")

    if report.errors:
        lines.append("## Errors")
        lines.append("")
        for error in report.errors:
            lines.append(f"- {error}")
        lines.append("")

    return "\n".join(lines)


def build_json_report(report: ValidationReport) -> Dict[str, Any]:
    return {
        "summary": {
            "errors": len(report.errors),
            "warnings": len(report.warnings),
            "status": "FAILED" if report.errors else "PASSED",
            "areas": [
                {
                    "area": area,
                    "result": result,
                    "checks": int(checks),
                    "ok": int(ok_count),
                    "needs_attention": int(issue_count),
                }
                for area, result, checks, ok_count, issue_count in build_area_summary_rows(report)
            ],
        },
        "infos": report.infos,
        "warnings": report.warnings,
        "errors": report.errors,
        "plain_english_summary": {
            table_name: {
                "result": summarize_rows(rows)["result"],
                "total_checks": summarize_rows(rows)["total_checks"],
                "ok_count": summarize_rows(rows)["ok_count"],
                "issue_count": summarize_rows(rows)["issue_count"],
                "issues": [build_issue_message(row) for row in summarize_rows(rows)["issue_rows"]],
            }
            for table_name, rows in report.tables.items()
        },
        "tables": {
            table_name: [
                {
                    "check": prettify_text(row.check),
                    "target": prettify_text(row.target),
                    "status": display_status(row.status),
                    "details": row.details,
                }
                for row in rows
            ]
            for table_name, rows in report.tables.items()
        },
    }


def ensure_output_dir(file_path: str) -> None:
    output_dir = os.path.dirname(file_path)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)


def write_markdown_report(report: ValidationReport, report_file: str) -> None:
    ensure_output_dir(report_file)
    with open(report_file, "w", encoding="utf-8") as file_handle:
        file_handle.write(build_markdown_report(report))


def write_json_report(report: ValidationReport, report_file: str) -> None:
    ensure_output_dir(report_file)
    with open(report_file, "w", encoding="utf-8") as file_handle:
        json.dump(build_json_report(report), file_handle, indent=2, ensure_ascii=False)


def write_word_report(report: ValidationReport, report_file: str) -> None:
    ensure_output_dir(report_file)
    document = Document()
    document.add_heading("Validation Report", level=0)

    summary = document.add_paragraph()
    summary.add_run("Status: ").bold = True
    summary.add_run("FAILED" if report.errors else "PASSED")
    summary.add_run("\nErrors: ").bold = True
    summary.add_run(str(len(report.errors)))
    summary.add_run("\nWarnings: ").bold = True
    summary.add_run(str(len(report.warnings)))

    document.add_heading("At A Glance", level=1)
    summary_table = document.add_table(rows=1, cols=5)
    summary_table.style = "Table Grid"
    summary_headers = ("Area", "Result", "Checks", "OK", "Needs attention")
    for index, header in enumerate(summary_headers):
        summary_table.rows[0].cells[index].text = header

    for row_values in build_area_summary_rows(report):
        row_cells = summary_table.add_row().cells
        for index, value in enumerate(row_values):
            row_cells[index].text = value

    for table_name, rows in report.tables.items():
        summary_data = summarize_rows(rows)
        document.add_heading(table_name, level=1)
        document.add_paragraph(f"Result: {summary_data['result']}")
        document.add_paragraph(f"Checks completed: {summary_data['total_checks']}")
        document.add_paragraph(f"Items OK: {summary_data['ok_count']}")
        document.add_paragraph(f"Items needing attention: {summary_data['issue_count']}")

        document.add_paragraph("What needs attention:")
        if summary_data["issue_rows"]:
            for row in summary_data["issue_rows"]:
                document.add_paragraph(build_issue_message(row), style="List Bullet")
        else:
            document.add_paragraph("No problems found in this section.", style="List Bullet")

        document.add_paragraph("Detailed check table:")
        table = document.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        header_cells = table.rows[0].cells
        word_headers = ("What was checked", "Item", "Result", "Explanation")
        for index, header in enumerate(word_headers):
            header_cells[index].text = header

        for row in rows:
            cells = table.add_row().cells
            cells[0].text = prettify_text(row.check)
            cells[1].text = prettify_text(row.target)
            cells[2].text = display_status(row.status)
            cells[3].text = row.details

    if report.warnings:
        document.add_heading("Warnings", level=1)
        for warning in report.warnings:
            document.add_paragraph(warning, style="List Bullet")

    if report.errors:
        document.add_heading("Errors", level=1)
        for error in report.errors:
            document.add_paragraph(error, style="List Bullet")

    document.save(report_file)


def print_report(report: ValidationReport) -> None:
    print()
    print("=== Validation Summary ===")
    for line in build_summary_table_lines(report):
        print(line)

    for table_name, rows in report.tables.items():
        summary = summarize_rows(rows)
        print()
        print(f"=== {table_name} ===")
        print(f"Result: {summary['result']}")
        print(f"Checks completed: {summary['total_checks']}")
        print(f"Items OK: {summary['ok_count']}")
        print(f"Items needing attention: {summary['issue_count']}")
        if summary["issue_rows"]:
            print("What needs attention:")
            for row in summary["issue_rows"]:
                print(f"- {build_issue_message(row)}")
        else:
            print("What needs attention:")
            print("- No problems found in this section.")

    for message in report.infos:
        print(f"INFO: {message}")
    for message in report.warnings:
        print(f"WARNING: {message}")
    for message in report.errors:
        print(f"ERROR: {message}")

    print()
    print(f"Validation finished with {len(report.errors)} error(s) and {len(report.warnings)} warning(s).")
    if report.errors:
        print("Status: FAILED")
    else:
        print("Status: PASSED")


def main() -> int:
    args = parse_args()
    report = ValidationReport()

    validate_docx(args.docx, report)
    json_paths = resolve_input_json_paths(args)
    mapping_workbook = choose_mapping_workbook(args)

    personal_data = load_json_file(json_paths["json"], report, "Personal details JSON", required=True)
    education_data = load_json_file(json_paths["education_json"], report, "Education advice JSON", required=False)
    socialcare_data = load_json_file(json_paths["socialcare_json"], report, "Social care advice JSON", required=False)
    health_data = load_json_file(json_paths["health_json"], report, "Health advice JSON", required=False)

    validate_personal_json(personal_data, report)
    validate_education_json(education_data, report)
    validate_socialcare_json(socialcare_data, report)
    validate_health_json(health_data, report)

    workbook = load_workbook_safe(mapping_workbook, report)
    if workbook is not None:
        validate_workbook_headers(workbook, report)
        validate_personal_mapping(workbook, personal_data, report)
        validate_section_mapping(workbook, EDUCATION_DETAILS_SHEET, education_data, report)
        validate_section_mapping(workbook, SOCIAL_CARE_DETAILS_SHEET, socialcare_data, report)
        validate_section_mapping(workbook, HEALTH_DETAILS_SHEET, health_data, report)

    if args.actual_output_docx and args.expected_output_docx:
        compare_docx_outputs(args.actual_output_docx, args.expected_output_docx, report)
    elif args.actual_output_docx or args.expected_output_docx:
        report.warn("DOCX output comparison runs only when both --actual-output-docx and --expected-output-docx are provided.")

    if args.report_file:
        write_markdown_report(report, args.report_file)
        report.info(f"Saved Markdown validation report: {args.report_file}")

    if args.json_report_file:
        write_json_report(report, args.json_report_file)
        report.info(f"Saved JSON validation report: {args.json_report_file}")

    if args.word_report_file:
        write_word_report(report, args.word_report_file)
        report.info(f"Saved Word validation report: {args.word_report_file}")

    print_report(report)
    return 0 if report.is_valid else 1


if __name__ == "__main__":
    sys.exit(main())