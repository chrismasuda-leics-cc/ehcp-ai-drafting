"""
API routes for EHCP document processing pipeline.
"""

import re as _re
from openai import AzureOpenAI
import os
import json
import re
import asyncio
import queue
from typing import List, Optional

from fastapi import APIRouter, Body, UploadFile, File, Header, HTTPException, Depends
from fastapi.responses import FileResponse, StreamingResponse

from app.settings import (
    TEMP_DIR, OUTPUT_DIR, TEMPLATE_DOCX, MAPPING_WORKBOOK, EXPECTED_OUTPUT_DOCX,
    USE_MANAGED_IDENTITY, AZURE_OPENAI_ENDPOINT, AZURE_OPENAI_API_KEY,
    AZURE_OPENAI_DEPLOYMENT, AZURE_OPENAI_API_VERSION,
    MODEL_MAX_TOKENS, get_openai_token_provider,
)
from app.models.schemas import (
    AnalyzeRequest, AnalyzeResponse, AnalyzeResult,
    WriteRequest, WriteResponse,
)
from app.services.orchestrator import run_maf_pipeline, run_maf_writer_pipeline
from app.services.blob_storage import (
    delete_blob, is_blob_storage_enabled,
    upload_file_to_blob, ensure_local_file, blob_key_for_path,
)
from app.services.audit_logger import log_action
from app.services.job_logger import (
    create_job_record, add_upload_document, add_upload_failure,
    set_analyse_start, set_analyse_complete, set_create_ehcp_complete,
    add_token_usage, add_completeness_entry, set_error, save_job_record,
    load_job_record,
)
from app.auth import get_current_user

router = APIRouter(prefix="/api", tags=["pipeline"])


# ---------------------------------------------------------
# Session-scoped directory helpers
# ---------------------------------------------------------

_SAFE_SESSION_RE = re.compile(r"^[a-zA-Z0-9_\-]{1,64}$")


def _session_temp_dir(session_id: str | None) -> str:
    """Return session-scoped temp dir.  Falls back to global TEMP_DIR."""
    if session_id and _SAFE_SESSION_RE.match(session_id):
        return os.path.join(TEMP_DIR, session_id)
    return TEMP_DIR


def _session_output_dir(session_id: str | None) -> str:
    """Return session-scoped output dir.  Falls back to global OUTPUT_DIR."""
    if session_id and _SAFE_SESSION_RE.match(session_id):
        return os.path.join(OUTPUT_DIR, session_id)
    return OUTPUT_DIR


async def _load_job_record(session_id: str | None, job_id: str | None) -> dict | None:
    """Load an existing job record from Cosmos DB, or return None."""
    if not job_id:
        return None
    return await load_job_record(job_id, session_id)


# Document type definitions
DOC_TYPE_MAP = {
    "Personal Details":   {"prompt": "prompts/personal_prompt.txt",   "schema": "schemas/personal_schema.json",   "key": "personal"},
    "Education Advice":   {"prompt": "prompts/education_prompt.txt",  "schema": "schemas/education_schema.json",  "key": "education"},
    "Health Advice":      {"prompt": "prompts/health_prompt.txt",     "schema": "schemas/health_schema.json",     "key": "health"},
    "Social Care Advice": {"prompt": "prompts/socialcare_prompt.txt", "schema": "schemas/socialcare_schema.json", "key": "socialcare"},
}

DOC_TYPE_OPTIONS = list(DOC_TYPE_MAP.keys())


def auto_detect_doc_type(filename: str) -> str:
    name_lower = filename.lower()
    if any(kw in name_lower for kw in ("education", "leps")):
        return "Education Advice"
    if "health" in name_lower:
        return "Health Advice"
    if any(kw in name_lower for kw in ("social", "_sc ", "_sc_", "_sc.")):
        return "Social Care Advice"
    stem = os.path.splitext(name_lower)[0]
    if stem.endswith("_sc") or "_sc " in name_lower or " sc " in name_lower or " sc_" in name_lower:
        return "Social Care Advice"
    return "Personal Details"


# ---------------------------------------------------------
# Content-based document type detection
# ---------------------------------------------------------

# Unique keyword markers per document type.  Each keyword contributes a
# score; the type with the highest total wins.  Keywords are checked
# case-insensitively against the first ~5000 chars of extracted text.

_CONTENT_MARKERS = {
    "Personal Details": [
        ("personal details", 3),
        ("ethnicity", 3),
        ("preferred name", 3),
        ("parent or carer first language", 3),
        ("main contact", 2),
        ("looked after child", 2),
        ("one page profile", 2),
        ("parent or carer 2", 2),
        ("method of communication", 2),
        ("child or young person first language", 2),
    ],
    "Education Advice": [
        ("education advice and information for statutory assessment", 4),
        ("educational psychology service", 3),
        ("specialist teaching services", 3),
        ("early years advisory teachers", 3),
        ("how i communicate", 3),
        ("educational psychology", 3),
        ("education advice", 3),
        ("cognition and learning", 3),
        ("communication and interaction", 3),
        ("sensory and physical", 3),
        ("journey so far", 2),
        ("pupil views", 2),
        ("social emotional mental health", 2),
        ("education setting", 2),
    ],
    "Health Advice": [
        ("contributions from health", 3),
        ("health advice", 3),
        ("medical history", 3),
        ("health needs", 3),
        ("health provision", 3),
        ("health and care plan assessment", 2),
        ("contributions requested", 2),
        ("nhs number", 2),
    ],
    "Social Care Advice": [
        ("social care advice and information for education health and care needs assessment", 4),
        ("social care advice", 3),
        ("social care needs", 3),
        ("social care provision", 3),
        ("mosaic id", 3),
        ("h1 social care", 2),
        ("h2 social care", 2),
        ("lac, cp, cin", 2),
        ("post-adoption", 2),
        ("post adoption", 2),
    ],
}


def _extract_text_quick(file_path: str, max_chars: int = 5000) -> str:
    """Extract the first `max_chars` of text from a DOCX or PDF without
    calling Azure Document Intelligence (fast, local-only)."""
    ext = os.path.splitext(file_path)[1].lower()
    text = ""
    try:
        if ext == ".docx":
            from docx import Document
            doc = Document(file_path)
            parts = []
            # Extract text from body paragraphs
            for para in doc.paragraphs:
                parts.append(para.text)
                if sum(len(p) for p in parts) >= max_chars:
                    break
            # Also extract text from table cells (Name, DOB etc. are often in tables)
            if sum(len(p) for p in parts) < max_chars:
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                parts.append(cell_text)
                            if sum(len(p) for p in parts) >= max_chars:
                                break
                        if sum(len(p) for p in parts) >= max_chars:
                            break
                    if sum(len(p) for p in parts) >= max_chars:
                        break
            text = "\n".join(parts)
        elif ext == ".pdf":
            import fitz
            pdf = fitz.open(file_path)
            parts = []
            for page in pdf:
                parts.append(page.get_text())
                if sum(len(p) for p in parts) >= max_chars:
                    break
            pdf.close()
            text = "\n".join(parts)
    except Exception:
        pass
    return text[:max_chars]


def detect_doc_type_by_content(file_path: str) -> str:
    """Detect document type by scoring unique content markers found in the
    document text.  Falls back to filename-based detection if no clear
    winner emerges."""
    text = _extract_text_quick(file_path).lower()
    if not text.strip():
        return auto_detect_doc_type(os.path.basename(file_path))

    scores = {}
    for doc_type, markers in _CONTENT_MARKERS.items():
        score = sum(weight for keyword, weight in markers if keyword in text)
        scores[doc_type] = score

    best_type = max(scores, key=scores.get)
    # Require a minimum score to trust content detection
    if scores[best_type] >= 4:
        return best_type
    # Fallback to filename-based detection
    return auto_detect_doc_type(os.path.basename(file_path))


# ---------------------------------------------------------
# Child name extraction from document content (via Azure OpenAI)
# ---------------------------------------------------------


_aoai_client = None


def _get_aoai_client() -> AzureOpenAI:
    global _aoai_client
    if _aoai_client is None:
        if USE_MANAGED_IDENTITY:
            _aoai_client = AzureOpenAI(
                azure_endpoint=AZURE_OPENAI_ENDPOINT,
                azure_ad_token_provider=get_openai_token_provider(),
                api_version=AZURE_OPENAI_API_VERSION,
            )
        else:
            _aoai_client = AzureOpenAI(
                azure_endpoint=AZURE_OPENAI_ENDPOINT,
                api_key=AZURE_OPENAI_API_KEY,
                api_version=AZURE_OPENAI_API_VERSION,
            )
    return _aoai_client


def _resolve_json_value(data: dict, json_path: str):
    """Resolve a dot-separated JSON path to a value in nested data."""
    current = data
    for part in json_path.split("."):
        if isinstance(current, dict):
            current = current.get(part)
        else:
            return None
        if current is None:
            return None
    return current


def _count_mapping_fields(output_data: dict, mapping_fields: list) -> tuple:
    """Count filled/empty fields based on mapping definitions.
    Returns (filled_list, empty_list) of json_path strings."""
    filled = []
    empty = []
    seen = set()
    for field in mapping_fields:
        jp = field["json_path"]
        if jp in seen:
            continue
        seen.add(jp)
        val = _resolve_json_value(output_data, jp)
        if val is None:
            empty.append(jp)
        elif isinstance(val, list):
            if val and any(item for item in val if item):
                filled.append(jp)
            else:
                empty.append(jp)
        elif isinstance(val, dict):
            if any(v for v in val.values() if v):
                filled.append(jp)
            else:
                empty.append(jp)
        elif str(val).strip():
            filled.append(jp)
        else:
            empty.append(jp)
    return filled, empty


def _extract_name_from_filename(file_path: str) -> str | None:
    """Layer 2: Extract a plausible child name from the filename itself.
    Strips doc-type keywords, separators, and extensions, then returns
    tokens that look like name parts (capitalised words, 2+ chars)."""
    stem = os.path.splitext(os.path.basename(file_path))[0]

    # Replace underscores, hyphens, dots with spaces
    stem = _re.sub(r"[_\-\.]+", " ", stem)

    # Remove common doc-type keywords (case-insensitive)
    noise = (
        r"\b(health|education|social\s*care|personal|advice|report|details|"
        r"ehcp|leps|copy|sc|docx?|pdf|section|part|draft|final|v\d+|"
        r"child|student|pupil|young|person|boy|girl)\b"
    )
    cleaned = _re.sub(noise, "", stem, flags=_re.IGNORECASE).strip()

    # Keep only alphabetic tokens with 2+ chars (ignore initials, numbers)
    tokens = [t for t in cleaned.split() if len(t) >= 2 and t.isalpha()]
    if not tokens:
        return None

    # Capitalise each token as a proper name
    name = " ".join(t.capitalize() for t in tokens)
    print(
        f"[extract_name_from_filename] {os.path.basename(file_path)} -> '{name}'")
    return name


def extract_child_name(file_path: str) -> str | None:
    """Extract the child's full name using a strict 2-layer approach:

    Layer 1 — LLM: Send the first ~2000 chars of document text + filename
              to Azure OpenAI for name extraction.
    Layer 2 — Filename parsing: If Layer 1 fails, parse the filename itself
              to extract a plausible name from naming conventions like
              'Ruben_Amos_Health_Advice.docx'.
    """
    filename = os.path.basename(file_path)
    text = _extract_text_quick(file_path, max_chars=5000)

    # ── Layer 1: LLM-based extraction ──
    if text.strip():
        try:
            client = _get_aoai_client()
            _GENERIC_WORDS = {
                "child", "student", "pupil", "young", "person",
                "boy", "girl", "infant", "toddler", "teenager",
                "unknown", "n/a", "none", "the",
            }
            resp = client.chat.completions.create(
                model=AZURE_OPENAI_DEPLOYMENT,
                messages=[
                    {
                        "role": "system",
                        "content": (
                            "You are a data extraction assistant. "
                            "Extract the full name of the child or young person this document is about. "
                            "Use both the document content AND the filename as clues. "
                            "Return ONLY the full name including any middle names, "
                            "always in 'Firstname [Middlename] Lastname' order. "
                            "Do NOT return any other text, explanation, or punctuation. "
                            "Do NOT return generic words like 'Child', 'Student', 'Young Person'. "
                            "If you cannot determine the child's name, return exactly: UNKNOWN"
                        ),
                    },
                    {
                        "role": "user",
                        "content": f"Filename: {filename}\n\n{text}",
                    },
                ],
                temperature=0.0,
                seed=42,
                max_completion_tokens=MODEL_MAX_TOKENS,
                timeout=30,
            )
            name = resp.choices[0].message.content.strip()
            print(
                f"[extract_child_name] Layer 1 (LLM): {filename} -> '{name}'")
            # Reject generic/non-name words
            if name and name.upper() != "UNKNOWN" and len(name) >= 2:
                name_tokens = set(name.lower().split())
                if not name_tokens.issubset(_GENERIC_WORDS):
                    return name
        except Exception as e:
            print(f"[extract_child_name] Layer 1 ERROR for {filename}: {e}")

    # ── Layer 2: Filename-based extraction ──
    fallback = _extract_name_from_filename(file_path)
    if fallback:
        print(
            f"[extract_child_name] Layer 2 (filename): {filename} -> '{fallback}'")
        return fallback

    print(f"[extract_child_name] Both layers failed for {filename}")
    return None


@router.get("/health")
async def health_check():
    return {"status": "ok"}


@router.get("/doc-types")
async def get_doc_types():
    return {"doc_types": DOC_TYPE_OPTIONS}


# ---------------------------------------------------------
# Mapping fields endpoint — returns the output-document fields
# from the mapping Excel, per doc type.
# ---------------------------------------------------------

_mapping_fields_cache = None


def _load_mapping_fields() -> dict:
    """Load the mapping Excel and return a dict of doc_type -> list of
    {json_path, label} entries representing the fields that appear in the
    final EHCP output document."""
    global _mapping_fields_cache
    if _mapping_fields_cache is not None:
        return _mapping_fields_cache

    import openpyxl
    wb = openpyxl.load_workbook(MAPPING_WORKBOOK, read_only=True)
    result = {}

    # --- PersonalDetails ---
    ws = wb["PersonalDetails"]
    personal_fields = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        json_key = row[3]  # column D: json_key
        label = row[2]     # column C: label
        if json_key:
            personal_fields.append(
                {"json_path": json_key, "label": label or json_key})
    result["personal"] = personal_fields

    # --- EducationAdvice, HealthAdvice, SocialCareAdvice ---
    for sheet_name, doc_key in [
        ("EducationAdvice", "education"),
        ("HealthAdvice", "health"),
        ("SocialCareAdvice", "socialcare"),
    ]:
        ws = wb[sheet_name]
        fields = []
        for row in ws.iter_rows(min_row=2, values_only=True):
            key = row[2]      # column C: key
            section = row[3]  # column D: section
            label = row[0]    # column A: Template column (display label)
            if not key:
                continue
            # Some keys have comma-separated values like "How_I_communicate , Other_things_Id_like_people_to_know"
            for k in key.split(","):
                k = k.strip()
                if not k:
                    continue
                json_path = f"{section}.{k}" if section else k
                fields.append({"json_path": json_path, "label": label or k})
        result[doc_key] = fields

    wb.close()
    _mapping_fields_cache = result
    return result


@router.get("/mapping-fields")
async def get_mapping_fields():
    """Return the output-document field definitions from the mapping Excel."""
    return _load_mapping_fields()


@router.post("/log-browse")
async def log_browse_files(
    filenames: List[str] = Body(default=[]),
    x_session_id: str | None = Header(None),
    current_user: Optional[dict] = Depends(get_current_user),
):
    """Log when a user selects/browses files in the file picker."""
    await log_action(
        action="browse_file",
        session_id=x_session_id,
        user=current_user,
        details={"files": filenames},
    )
    return {"logged": True}


@router.post("/log-activity")
async def log_activity(
    payload: dict = Body(...),
    x_session_id: str | None = Header(None),
    current_user: Optional[dict] = Depends(get_current_user),
):
    """Generic audit log endpoint for frontend UI actions."""
    action = payload.get("action", "unknown")
    details = payload.get("details", {})
    await log_action(
        action=action,
        session_id=x_session_id,
        user=current_user,
        details=details,
    )
    return {"logged": True}


@router.post("/upload")
async def upload_files(
    files: List[UploadFile] = File(...),
    x_session_id: str | None = Header(None),
    x_job_id: str | None = Header(None),
    current_user: Optional[dict] = Depends(get_current_user),
):
    """Upload files to the backend temp directory."""
    temp_dir = _session_temp_dir(x_session_id)
    os.makedirs(temp_dir, exist_ok=True)

    # Reuse existing job record if job_id provided, otherwise create new.
    # When creating, honour the caller-supplied job_id so repeated uploads for
    # the same case map to a single record (id == job_id) instead of creating
    # duplicate job records.
    job = None
    if x_job_id:
        job = await _load_job_record(x_session_id, x_job_id)
    if not job:
        job = create_job_record(
            user=current_user, session_id=x_session_id, job_id=x_job_id)

    uploaded = []
    current_failures = []
    for file in files:
        file_path = os.path.join(temp_dir, file.filename)
        content = await file.read()
        file_size = len(content)

        try:
            with open(file_path, "wb") as f:
                f.write(content)

            # Upload to blob storage using a session-scoped key so that a later
            # analyze/write step running on a different replica can retrieve it.
            if is_blob_storage_enabled():
                await asyncio.to_thread(upload_file_to_blob, file_path)

            # Detect doc type by content first, fall back to filename.
            # These do blocking work (docx parsing + a synchronous Azure
            # OpenAI call), so run them off the event loop; otherwise one slow
            # upload stalls every other concurrent upload on the same worker
            # and causes one user's upload to time out.
            detected = await asyncio.to_thread(
                detect_doc_type_by_content, file_path)

            # Extract child name from document content (blocking LLM call)
            child_name = await asyncio.to_thread(extract_child_name, file_path)

            uploaded.append({
                "filename": file.filename,
                "path": file_path,
                "detected_type": detected,
                "child_name": child_name,
            })

            # Add to job record
            add_upload_document(job, file.filename, file_size, detected)

        except Exception as e:
            add_upload_failure(job, file.filename, str(e))
            current_failures.append(file.filename)

    # Persist initial job record so later pipeline stages can load it
    await save_job_record(job)

    # Audit log (existing)
    await log_action(
        action="upload",
        session_id=x_session_id,
        user=current_user,
        details={"files": [{"filename": u["filename"],
                            "doc_type": u["detected_type"]} for u in uploaded],
                 "job_id": job["job_id"]},
    )

    if current_failures:
        failed_names = ", ".join(current_failures)
        raise HTTPException(
            status_code=503,
            detail=f"Upload failed for {failed_names}. Please retry.",
        )

    return {"uploaded": uploaded, "job_id": job["job_id"]}


@router.post("/analyze", response_model=AnalyzeResponse)
async def analyze_documents(
    request: AnalyzeRequest,
    x_session_id: str | None = Header(None),
    current_user: Optional[dict] = Depends(get_current_user),
):
    """Run the reader pipeline (Reader → Extractor → Validator → QualityChecker)."""
    temp_dir = _session_temp_dir(x_session_id)
    os.makedirs(temp_dir, exist_ok=True)

    file_configs = []
    for fc in request.files:
        base_name = os.path.splitext(fc.filename)[0]
        doc_type = fc.doc_type if fc.doc_type in DOC_TYPE_MAP else auto_detect_doc_type(
            fc.filename)
        info = DOC_TYPE_MAP[doc_type]

        file_configs.append({
            "input_docx": os.path.join(temp_dir, fc.filename),
            "prompt_file": info["prompt"],
            "schema_file": info["schema"],
            "output_file": os.path.join(temp_dir, f"{base_name}_output.json"),
            "validation_output_file": os.path.join(temp_dir, f"{base_name}_validation.json"),
        })

    # Inputs may have been uploaded on a different replica — pull them from
    # blob storage if they are missing locally.
    for cfg in file_configs:
        ensure_local_file(cfg["input_docx"])

    # Run the multi-agent pipeline
    maf_results, document_texts = await run_maf_pipeline(file_configs)

    results = []
    for i, fc in enumerate(request.files):
        base_name = os.path.splitext(fc.filename)[0]
        if isinstance(maf_results[i], Exception):
            results.append(AnalyzeResult(
                filename=fc.filename,
                output_file="",
                validation_file="",
                success=False,
                error=str(maf_results[i]),
            ))
        else:
            results.append(AnalyzeResult(
                filename=fc.filename,
                output_file=f"{base_name}_output.json",
                validation_file=f"{base_name}_validation.json",
                success=True,
            ))

    # Audit log
    await log_action(
        action="analyze",
        session_id=x_session_id,
        user=current_user,
        details={
            "files": [{"filename": r.filename, "success": r.success, "error": r.error} for r in results],
        },
    )

    return AnalyzeResponse(results=results)


@router.post("/analyze-stream")
async def analyze_documents_stream(
    request: AnalyzeRequest,
    x_session_id: str | None = Header(None),
    current_user: Optional[dict] = Depends(get_current_user),
):
    """Run the reader pipeline with Server-Sent Events for progress updates."""
    # Load or create job record for this session
    job_id = request.job_id if request.job_id else None

    # Log analyze start immediately
    await log_action(
        action="analyze_start",
        session_id=x_session_id,
        user=current_user,
        details={"files": [{"filename": fc.filename,
                            "doc_type": fc.doc_type} for fc in request.files],
                 "job_id": job_id},
    )

    temp_dir = _session_temp_dir(x_session_id)
    os.makedirs(temp_dir, exist_ok=True)

    file_configs = []
    for fc in request.files:
        base_name = os.path.splitext(fc.filename)[0]
        doc_type = fc.doc_type if fc.doc_type in DOC_TYPE_MAP else auto_detect_doc_type(
            fc.filename)
        info = DOC_TYPE_MAP[doc_type]
        file_configs.append({
            "input_docx": os.path.join(temp_dir, fc.filename),
            "prompt_file": info["prompt"],
            "schema_file": info["schema"],
            "output_file": os.path.join(temp_dir, f"{base_name}_output.json"),
            "validation_output_file": os.path.join(temp_dir, f"{base_name}_validation.json"),
            "doc_type": doc_type,
        })

    # Inputs may have been uploaded on a different replica — pull them from
    # blob storage if they are missing locally.
    for cfg in file_configs:
        ensure_local_file(cfg["input_docx"])

    progress_queue = queue.Queue()

    def progress_callback(event):
        progress_queue.put(event)

    # Update job record with analyse start (saved only at end of pipeline)
    job = await _load_job_record(x_session_id, job_id)
    if job:
        set_analyse_start(job)

    async def event_generator():
        nonlocal job
        # Run pipeline in a background thread so we can stream progress
        loop = asyncio.get_event_loop()

        async def run_pipeline():
            return await run_maf_pipeline(file_configs, progress_callback=progress_callback)

        task = asyncio.create_task(run_pipeline())

        # Stream progress events until pipeline finishes
        while not task.done():
            try:
                event = progress_queue.get_nowait()
                yield f"data: {json.dumps(event)}\n\n"
            except queue.Empty:
                await asyncio.sleep(0.3)

        # Drain remaining events
        while not progress_queue.empty():
            event = progress_queue.get_nowait()
            yield f"data: {json.dumps(event)}\n\n"

        # Send final results
        try:
            maf_results, document_texts = task.result()
            results = []
            for i, fc in enumerate(request.files):
                base_name = os.path.splitext(fc.filename)[0]
                if isinstance(maf_results[i], Exception):
                    results.append({"filename": fc.filename, "output_file": "",
                                   "validation_file": "", "success": False, "error": str(maf_results[i])})
                else:
                    results.append({"filename": fc.filename, "output_file": f"{base_name}_output.json",
                                   "validation_file": f"{base_name}_validation.json", "success": True})

            # Update job record with analyse complete + completeness
            if job:
                has_error = any(not r["success"] for r in results)
                error_msg = "; ".join(r.get("error", "")
                                      for r in results if not r["success"])
                set_analyse_complete(job, error=error_msg if has_error else "")

                # Add completeness report per document (using mapping fields, same as frontend)
                mapping_fields = _load_mapping_fields()
                for i, fc_cfg in enumerate(file_configs):
                    if not isinstance(maf_results[i], Exception):
                        output_path = fc_cfg["output_file"]
                        if os.path.exists(output_path):
                            try:
                                with open(output_path, "r", encoding="utf-8") as of:
                                    output_data = json.load(of)
                                doc_key = DOC_TYPE_MAP.get(
                                    fc_cfg["doc_type"], {}).get("key", "")
                                fields_def = mapping_fields.get(doc_key, [])
                                filled, empty = _count_mapping_fields(
                                    output_data, fields_def)
                                total = len(filled) + len(empty)
                                source_file = os.path.basename(
                                    fc_cfg["input_docx"])
                                add_completeness_entry(
                                    job, fc_cfg["doc_type"],
                                    total, len(filled), len(empty), empty,
                                    source_file=source_file
                                )
                            except Exception:
                                pass

                # Accumulate token usage from all documents
                for i, _ in enumerate(file_configs):
                    if not isinstance(maf_results[i], Exception):
                        usage = maf_results[i].get("token_usage", {})
                        if usage:
                            add_token_usage(
                                job,
                                prompt_tokens=usage.get("prompt_tokens", 0),
                                completion_tokens=usage.get(
                                    "completion_tokens", 0),
                                total_tokens=usage.get("total_tokens", 0),
                            )

                await save_job_record(job)

            # Audit log after pipeline completes
            await log_action(
                action="analyze_complete",
                session_id=x_session_id,
                user=current_user,
                details={"files": [{"filename": r["filename"], "success": r["success"], "error": r.get(
                    "error")} for r in results],
                    "job_id": job_id},
            )

            yield f"data: {json.dumps({'type': 'complete', 'results': results})}\n\n"
        except Exception as e:
            if job:
                set_error(job, str(e))
                await save_job_record(job)
            yield f"data: {json.dumps({'type': 'error', 'error': str(e)})}\n\n"

    return StreamingResponse(event_generator(), media_type="text/event-stream")


@router.post("/write-ehcp", response_model=WriteResponse)
async def write_ehcp(
    request: WriteRequest,
    x_session_id: str | None = Header(None),
    current_user: Optional[dict] = Depends(get_current_user),
):
    """Run the writer pipeline (TemplateWriter → WriterValidator)."""
    if not os.path.exists(TEMPLATE_DOCX):
        raise HTTPException(status_code=400, detail="Template DOCX not found")
    if not os.path.exists(MAPPING_WORKBOOK):
        raise HTTPException(
            status_code=400, detail="Mapping workbook not found")

    temp_dir = _session_temp_dir(x_session_id)

    # Resolve json paths — frontend sends filenames, resolve to temp/ paths.
    # The extraction step may have run on a different replica, so pull each
    # JSON from blob storage if it is missing on this one.
    json_paths = {}
    for key, path in request.json_paths.items():
        if not path:
            json_paths[key] = None
            continue
        if os.path.isabs(path):
            candidate = path
        else:
            candidate = os.path.join(temp_dir, os.path.basename(path))
        ensure_local_file(candidate)
        if os.path.exists(candidate):
            json_paths[key] = os.path.abspath(candidate)
        else:
            json_paths[key] = None

    result = await run_maf_writer_pipeline(
        template_docx=TEMPLATE_DOCX,
        output_docx=None,
        json_paths=json_paths,
        mapping_workbook=MAPPING_WORKBOOK,
        expected_output_docx=EXPECTED_OUTPUT_DOCX,
        output_dir_override=_session_output_dir(x_session_id),
        report_dir_override=_session_temp_dir(x_session_id),
    )

    # Return just the filename so the frontend can use /results/{filename}
    filled_path = result.get("filled_docx_path")
    filled_filename = os.path.basename(filled_path) if filled_path else None
    filled_size = os.path.getsize(
        filled_path) if filled_path and os.path.exists(filled_path) else 0

    # Update job record with write-ehcp completion
    job = await _load_job_record(x_session_id, request.job_id)
    if job:
        storage_path = f"outputs/{filled_filename}" if filled_filename else ""
        set_create_ehcp_complete(
            job, filled_filename or "", filled_size, storage_path)
        job["status"] = "completed"
        await save_job_record(job)

    # Audit log
    await log_action(
        action="write_ehcp",
        session_id=x_session_id,
        user=current_user,
        details={"output_file": filled_filename, "json_paths": {
            k: os.path.basename(v) if v else None for k, v in json_paths.items()},
            "job_id": job["job_id"] if job else None},
    )

    return WriteResponse(
        filled_docx_path=filled_filename,
        validation_report=result.get("validation_report"),
    )


@router.get("/results/{filename}")
async def download_result(filename: str, x_session_id: str | None = Header(None)):
    """Download a result file for the caller's session only.

    The file is looked up in the session-scoped temp/output directories.  If it
    is not present on this replica (it may have been produced on another), it is
    retrieved from blob storage using the session-scoped key.  The global
    temp/output directories are intentionally NOT searched, to avoid serving a
    different user's file.
    """
    safe_name = os.path.basename(filename)
    temp_dir = _session_temp_dir(x_session_id)
    output_dir = _session_output_dir(x_session_id)

    for candidate in (
        os.path.join(temp_dir, safe_name),
        os.path.join(output_dir, safe_name),
    ):
        if ensure_local_file(candidate):
            return FileResponse(candidate, filename=safe_name)

    raise HTTPException(status_code=404, detail="File not found")


@router.get("/download/{filepath:path}")
async def download_file(filepath: str):
    """Download any file by relative path."""
    if not os.path.exists(filepath):
        raise HTTPException(status_code=404, detail="File not found")
    return FileResponse(filepath, filename=os.path.basename(filepath))


@router.delete("/files/{filename}")
async def delete_file(
    filename: str,
    x_session_id: str | None = Header(None),
    current_user: Optional[dict] = Depends(get_current_user),
):
    """Delete a file from temp directory and blob storage."""
    temp_dir = _session_temp_dir(x_session_id)
    file_path = os.path.join(temp_dir, os.path.basename(filename))
    if os.path.exists(file_path):
        os.remove(file_path)

    if is_blob_storage_enabled():
        try:
            delete_blob(blob_key_for_path(file_path))
        except Exception:
            pass

    # Audit log
    await log_action(
        action="delete",
        session_id=x_session_id,
        user=current_user,
        details={"filename": filename},
    )

    return {"deleted": filename}
